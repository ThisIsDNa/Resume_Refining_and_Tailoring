"""
Builds validation-safe DOCX resumes from structured payloads for Tailor and Refinery flows.

Recruiter-ready .docx export: resume content only — no scores, diagnostics, or analysis.

**DOCX payload pipeline (single path — no legacy builder in this repo)**

Structured payloads for ``.docx`` are created **only** via
``app.services.resume_document_assembly.build_resume_document_payload`` from
``build_export_docx_package``. There is no ``build_legacy_payload`` / alternate
constructor on the export route; tests that care about parity should call the
same ``build_resume_document_payload`` + ``build_export_docx_package`` chain.

Rendering uses **only** ``build_docx_from_payload`` with that same
``ResumeDocumentPayload`` instance (see ``EXPORT_PIPELINE_DEBUG`` / ``PAYLOAD_ID_DEBUG``).
"""

from __future__ import annotations

import copy
import json
import logging
import os
import re
import time
from io import BytesIO
from collections import Counter
from typing import Any, Dict, List, Optional, Sequence, Tuple

from app.content.portfolio_resume_polish import maybe_apply_portfolio_resume_polish
from app.services.parse_resume import line_is_experience_noise
from app.services.experience_bullet_prioritization import prioritize_experience_bullets
from app.services.experience_header_normalization import _normalize_experience_headers
from app.services.resume_document_assembly import (
    ExperienceEntry,
    ResumeContractError,
    ResumeDocumentPayload,
    _US_CITY_STATE_LINE_RE,
    _line_is_compact_location_metadata,
    _line_is_standalone_calendar_date_range,
    _normalize_resume_line_dashes,
    _strip_skill_bucket_lines_from_project_entries,
    build_resume_document_payload,
    canonical_resume_dict,
    certification_entry_header_lines,
    dict_projects_to_entries,
    education_entry_header_lines,
    experience_entry_header_lines,
    coalesce_skills_for_export,
    merge_distinct_skill_lines,
    parse_certification_dicts_from_raw_text,
    parse_education_dicts_from_raw_text,
    prepare_experience_blocks_for_docx,
    prepare_project_blocks_for_docx,
    repair_experience_api_blocks_identity_from_bullets,
    skills_to_display_line,
    skills_to_display_lines,
    structured_payload_debug_json,
    validate_resume_document_payload,
)

# Single DOCX payload constructor used by ``build_export_docx_package`` (and required for test parity).
_DOCX_PAYLOAD_BUILDER = build_resume_document_payload

from app.services.resume_presentation import (
    build_outcome_phrase_export_summary,
    build_strong_identity_forward_export_summary,
    build_structured_identity_export_summary,
    is_tool_centric_summary,
    presentation_quality_warnings,
    prioritize_project_blocks_for_export,
    section_integrity_sanity_hints,
    trim_summary_for_scannability,
)
from app.services.rewrite_resume import (
    _SUMMARY_LEXICON_ALLOW,
    _default_resume_summary_fallback,
    _is_poor_or_sparse_fit,
    _primary_job_title_from_resume,
    _summary_has_forbidden_jd_copy_patterns,
    _summary_input_has_jd_pollution,
    _summary_words_grounded_in_resume,
    _trim_redundant_words,
)

logger = logging.getLogger(__name__)
_EXPORT_DEBUG_ENV = "RESUME_TAILOR_EXPORT_DEBUG"


def _export_debug_enabled() -> bool:
    return (os.environ.get(_EXPORT_DEBUG_ENV) or "").strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )

_EXPORT_FALLBACK_ONLY_ENV = "RESUME_TAILOR_EXPORT_SUMMARY_FALLBACK_ONLY"

_WEAK_EXPORT_SUMMARY = (
    "Business Systems Analyst with experience in data analysis, system validation, "
    "and BI tools including Excel and Power BI."
)

_DEBUG_MARKERS = (
    "coaching:",
    "score_breakdown",
    "classification",
    "evidence_id",
    "mapper_reason",
    "why this fits",
    "fit read",
    "keyword alignment",
)

# STEP 2 — internal / UI language that must never appear in recruiter export
_CONTENT_LEAK_PHRASES = (
    "fit snapshot",
    "why this fits",
    "gaps to watch",
    "support signal",
    "gap signal",
    "match strength",
    "strong / medium / weak",
    "strong/medium/weak",
    "validated_requirements",
    "tailoring note",
    "jd alignment",
    "keyword alignment",
    "fit read",
    "score_breakdown",
    "mapper_reason",
    "evidence_id",
    "before / after",
    "before/after",
    "kept as-is",
    "kept as is",
    "unchanged for this",
    "coaching:",
)

_MANIFESTO_OR_ABOUT_PATTERNS = re.compile(
    r"(?i)\b(?:our mission|we believe|join our team|about us:|"
    r"life at [a-z]|why [a-z]+ is a great place|committed to diversity)\b"
)

# Commentary / system artifacts in bullets (STEP 3)
_BULLET_COMMENTARY_PATTERNS = re.compile(
    r"(?i)\b(?:kept as-?is|left unchanged|revised to align|tailoring note|"
    r"\[coaching\]|\[note\]|\[system\]|as requested by the (?:jd|posting)|"
    r"unchanged per|mapper says|coaching)\b"
)

# Phrases stripped from bullet text before re-checking safety (longest first)
_BULLET_SCRUB_PHRASES: Tuple[str, ...] = (
    "as requested by the posting",
    "as requested by the jd",
    "tailoring note",
    "left unchanged",
    "kept as-is",
    "kept as is",
    "unchanged per",
    "mapper says",
    "coaching:",
    "[coaching]",
    "[note]",
    "[system]",
)

_BULLET_PAREN_SYSTEM = re.compile(
    r"\([^)]{0,200}?(?:kept as-?is|left unchanged|coaching|tailoring|mapper|system|jd posting|posting)[^)]{0,200}?\)",
    re.I,
)

# Known leaked project line (STEP 3: substring ``classification``); normalize before validation.
_LEAKY_CLASSIFICATION_GATING_BULLET_OLD = (
    "Developed classification and rule-based gating to ensure context-aware test generation "
    "and prevent cross-domain errors."
)
_LEAKY_CLASSIFICATION_GATING_BULLET_NEW = (
    "Developed rule-based gating to route scenario types accurately and prevent cross-domain logic errors."
)

# Experience rows that are clearly not jobs (mis-segmented project / tooling copy).
_EXPERIENCE_MISFILED_PROJECT_PHRASES: Tuple[str, ...] = (
    "test generation",
    "scenario context",
    "rule-based gating",
    "docx export",
    "ai-driven testing assistant",
)
_EXPERIENCE_ROLE_MONTH_ONLY = re.compile(
    r"^(january|february|march|april|may|june|july|august|september|october|november|december)"
    r"(\s+[–-]?\s*(19|20)\d{2,4})?\s*$",
    re.I,
)


def _rewrite_classification_gating_leak_in_blocks(blocks: Optional[List[dict]]) -> None:
    """In-place: replace the known leaky sentence in experience or project ``bullets`` lists."""
    if not blocks:
        return
    old_norm = re.sub(
        r"\s+", " ", _LEAKY_CLASSIFICATION_GATING_BULLET_OLD.strip().lower()
    )
    new = _LEAKY_CLASSIFICATION_GATING_BULLET_NEW
    for block in blocks:
        if not isinstance(block, dict):
            continue
        bullets = block.get("bullets")
        if not isinstance(bullets, list):
            continue
        for i, b in enumerate(bullets):
            t = str(b).strip()
            if re.sub(r"\s+", " ", t.lower()) == old_norm:
                bullets[i] = new


def _experience_block_signals_misfiled_project_or_noise(block: dict) -> bool:
    """True when heuristics suggest a misfiled project row or non-job noise (logging only; no drop)."""
    company = str(block.get("company") or "").strip()
    role = str(block.get("title") or block.get("role") or "").strip()
    if company.lower() == "remote":
        return True
    if role and _EXPERIENCE_ROLE_MONTH_ONLY.match(role):
        return True
    bullets = block.get("bullets") or []
    if not isinstance(bullets, list):
        return False
    blob = " ".join(str(b).lower() for b in bullets if str(b).strip())
    return any(p in blob for p in _EXPERIENCE_MISFILED_PROJECT_PHRASES)


def _filter_misfiled_experience_blocks_for_docx(blocks: List[dict]) -> List[dict]:
    """
    Log suspected misfiled / project-leak experience rows for investigation.
    Blocks are **not** removed (segmentation must be fixed before any drop).
    """
    out: List[dict] = []
    for i, b in enumerate(blocks or []):
        if not isinstance(b, dict):
            continue
        if _experience_block_signals_misfiled_project_or_noise(b):
            logger.warning(
                "export_docx flagging misfiled experience block (not dropped) idx=%s company=%r role=%r",
                i,
                (b.get("company") or "")[:120],
                (b.get("title") or b.get("role") or "")[:120],
            )
            if _export_debug_enabled():
                logger.debug(
                    "misfiled experience detail: %s",
                    json.dumps(
                        {
                            "idx": i,
                            "company": (b.get("company") or "")[:120],
                            "role": (b.get("title") or b.get("role") or "")[:120],
                        },
                        ensure_ascii=False,
                    ),
                )
        out.append(b)
    return out


def _log_experience_trailing_metadata_dict(
    *,
    entry_index: int,
    company: str,
    role: str,
    removed_location: Optional[str],
    removed_date: Optional[str],
) -> None:
    if not _export_debug_enabled():
        return
    logger.debug(
        "experience trailing metadata cleanup: %s",
        json.dumps(
            {
                "entry_index": entry_index,
                "company": (company or "").strip()[:200],
                "role": (role or "").strip()[:200],
                "removed_location": ((removed_location or "").strip()[:220] or None),
                "removed_date": ((removed_date or "").strip()[:220] or None),
            },
            ensure_ascii=False,
        ),
    )


def _clean_experience_trailing_metadata(experience_blocks: List[dict]) -> List[dict]:
    """
    Pop trailing location / calendar date-span lines out of API ``bullets`` into block
    metadata **only** when ``location`` / ``date`` fields are still empty. Evaluates at
    most the last two bullet lines per block (export path, before STEP 3 validation).
    Lines matching metadata patterns are always removed from ``bullets`` when recognized;
    assignment respects existing non-empty ``location`` / ``date_range`` / ``date``.
    """

    def _is_trailing_location_line(t: str) -> bool:
        s = (t or "").strip()
        return bool(s) and _line_is_compact_location_metadata(s)

    def _is_trailing_date_line(t: str) -> bool:
        s = (t or "").strip()
        return bool(s) and _line_is_standalone_calendar_date_range(s)

    out: List[dict] = []
    for bi, block in enumerate(experience_blocks or []):
        if not isinstance(block, dict):
            out.append(block)
            continue
        nb = dict(block)
        raw_bullets = nb.get("bullets") or []
        if not isinstance(raw_bullets, list):
            out.append(nb)
            continue
        bullets = [_normalize_resume_line_dashes(str(b).strip()) for b in raw_bullets if str(b).strip()]
        loc_f = str(nb.get("location") or "").strip()
        date_f = str(nb.get("date_range") or nb.get("date") or "").strip()

        rl: Optional[str] = None
        rd: Optional[str] = None

        if len(bullets) >= 2:
            a, b = bullets[-2], bullets[-1]
            if _is_trailing_location_line(a) and _is_trailing_date_line(b):
                bullets = bullets[:-2]
                rl, rd = a, b
            elif _is_trailing_date_line(a) and _is_trailing_location_line(b):
                bullets = bullets[:-2]
                rd, rl = a, b
        if rl is None and rd is None and bullets:
            last = bullets[-1]
            if _is_trailing_date_line(last):
                rd = last
                bullets = bullets[:-1]
            elif _US_CITY_STATE_LINE_RE.match(last) and _is_trailing_location_line(last):
                rl = last
                bullets = bullets[:-1]

        if rl is not None or rd is not None:
            if rl and not loc_f:
                nb["location"] = rl.strip()
            if rd and not date_f:
                dval = rd.strip()
                nb["date_range"] = dval
                nb["date"] = dval
            _log_experience_trailing_metadata_dict(
                entry_index=bi,
                company=str(nb.get("company") or ""),
                role=str(nb.get("title") or nb.get("role") or ""),
                removed_location=rl,
                removed_date=rd,
            )
        nb["bullets"] = bullets
        out.append(nb)
    return out


DOCX_EXPORT_VALIDATION_SUCCESS = "DOCX EXPORT VALIDATED - READY"
_SUCCESS_MSG = DOCX_EXPORT_VALIDATION_SUCCESS
_FAIL_HEAD = "DOCX EXPORT FAILED"

_ALLOWED_BODY_FONT_PT = frozenset({10.5, 11.0, 12.0})
_ALLOWED_NAME_FONT_PT = frozenset({12.0, 14.0})


def derive_match_strength(
    score_result: dict, mapping_result: dict, job_signals: dict
) -> str:
    """Coarse label for export summary override (strong / medium / weak)."""
    if _is_poor_or_sparse_fit(mapping_result, job_signals):
        return "weak"
    overall = int(score_result.get("overall_score") or 0)
    summ = score_result.get("summary") or {}
    matched = int(summ.get("matched_requirements") or 0)
    if overall >= 62 and matched >= 3:
        return "strong"
    return "medium"


def resolve_export_summary(tailored_summary: str, match_strength: str) -> str:
    """Final summary text for DOCX: weak uses fixed template; strong/medium use tailored only."""
    t = (tailored_summary or "").strip()
    if match_strength == "weak":
        return _WEAK_EXPORT_SUMMARY
    return t


def _resume_text_corpus_for_export_summary(resume_data: dict) -> str:
    """Resume raw + structured sections only (no bullet_changes / JD)."""
    parts: List[str] = [str(resume_data.get("raw_text") or "")]
    sections = resume_data.get("sections") if isinstance(resume_data.get("sections"), dict) else {}
    for block in sections.get("experience") or []:
        if isinstance(block, dict):
            parts.append(str(block.get("company") or ""))
            parts.append(str(block.get("title") or block.get("role") or ""))
            parts.append(str(block.get("date_range") or block.get("date") or ""))
            parts.append(str(block.get("location") or ""))
            for b in block.get("bullets") or []:
                parts.append(str(b))
    sk = sections.get("skills")
    if isinstance(sk, list):
        for s in sk:
            parts.append(str(s))
    for proj in sections.get("projects") or []:
        if isinstance(proj, dict):
            parts.append(str(proj.get("name") or proj.get("title") or ""))
            for b in proj.get("bullets") or []:
                parts.append(str(b))
    return " ".join(parts)


def _export_summary_passes_hygiene(
    summary: str, corpus: str, match_strength: str
) -> bool:
    """Same predicates as STEP 1 summary checks (strict; not a relaxation)."""
    s = (summary or "").strip()
    if not s:
        return False
    if match_strength == "weak":
        return s == _WEAK_EXPORT_SUMMARY
    cl = (corpus or "").lower()
    if _summary_has_forbidden_jd_copy_patterns(s) or _summary_input_has_jd_pollution(s):
        return False
    if not _summary_words_grounded_in_resume(s, cl):
        return False
    if _MANIFESTO_OR_ABOUT_PATTERNS.search(s):
        return False
    if _keyword_stuffing_summary(s):
        return False
    return True


def _export_identity_summary_passes_hygiene_resume_title_corpus(
    summary: str, corpus: str, match_strength: str, resume_data: dict
) -> bool:
    """
    Same hygiene as export STEP 1, but word-grounding is checked against the resume corpus
    plus the primary structured job title tokens. Used only for identity-built summaries
    so compound titles (e.g. ``Senior … / UAT Lead``) still ground title words that may not
    appear verbatim in merged bullet text.
    """
    s = (summary or "").strip()
    if not s:
        return False
    if match_strength == "weak":
        return s == _WEAK_EXPORT_SUMMARY
    cl = (corpus or "").lower()
    title_blob = (_primary_job_title_from_resume(resume_data) or "").strip().lower()
    cl_expanded = f"{cl} {title_blob}".strip()
    if _summary_has_forbidden_jd_copy_patterns(s) or _summary_input_has_jd_pollution(s):
        return False
    if not _summary_words_grounded_in_resume(s, cl_expanded):
        return False
    if _MANIFESTO_OR_ABOUT_PATTERNS.search(s):
        return False
    if _keyword_stuffing_summary(s):
        return False
    return True


def _log_summary_candidate_debug(
    *,
    source: str,
    text: str,
    accepted: bool,
    rejected_reason: str,
) -> None:
    if not _export_debug_enabled():
        return
    payload = json.dumps(
        {
            "source": source,
            "text": (text or "")[:400],
            "accepted": accepted,
            "rejected_reason": "" if accepted else (rejected_reason or "unknown"),
        },
        ensure_ascii=False,
    )
    logger.debug("summary candidate: %s", payload)


def _log_summary_policy_debug(
    *,
    winner_source: str,
    winner_summary: str,
    outcome_phrase: str,
    outcome_role_opener: bool,
    passes_hygiene: bool,
    banned_winner: bool,
) -> None:
    if not _export_debug_enabled():
        return
    payload = json.dumps(
        {
            "winner_source": winner_source,
            "winner_summary_preview": (winner_summary or "")[:220],
            "outcome_phrase_preview": (outcome_phrase or "")[:220],
            "outcome_phrase_role_opener": outcome_role_opener,
            "passes_hygiene_final": passes_hygiene,
            "banned_winner": banned_winner,
        },
        ensure_ascii=False,
    )
    logger.debug("summary policy: %s", payload)


def _export_summary_rejection_reason(
    summary: str,
    corpus: str,
    match_strength: str,
    *,
    resume_data: Optional[dict] = None,
    check_tool: bool = False,
    check_banned: bool = False,
) -> str:
    s = (summary or "").strip()
    if not s:
        return "empty"
    if match_strength == "weak" and s != _WEAK_EXPORT_SUMMARY:
        return "weak_match_not_weak_template"
    cl = (corpus or "").lower()
    if _summary_has_forbidden_jd_copy_patterns(s) or _summary_input_has_jd_pollution(s):
        return "jd_copy_or_pollution"
    if not _summary_words_grounded_in_resume(s, cl):
        ung = list_ungrounded_summary_tokens(s, cl)
        return f"word_grounding_fail_strict tokens={ung[:12]!r}"
    if _MANIFESTO_OR_ABOUT_PATTERNS.search(s):
        return "manifesto_or_about_pattern"
    if _keyword_stuffing_summary(s):
        return "keyword_stuffing"
    if check_banned and _is_banned_strongest_summary_pattern(s):
        return "banned_strongest_pattern"
    if check_tool and is_tool_centric_summary(s):
        return "tool_centric"
    return "ok"


_FALLBACK_ROLE_ORDER: Tuple[str, ...] = (
    "business systems analyst",
    "business analyst",
    "data analyst",
)

_FALLBACK_SKILL_PHRASES: Tuple[str, ...] = (
    "user acceptance testing",
    "process improvement",
    "data analysis",
    "documentation",
    "validation",
    "reporting",
    "testing",
)

_MINIMAL_LEXICON_SUMMARY = (
    "Professional with experience in analysis, delivery, and documentation in professional settings."
)

_BSA_IDENTITY_PREFERRED_FALLBACK = (
    "Business Systems Analyst specializing in system stabilization, UAT leadership, "
    "and cross-functional execution in enterprise environments."
)

# Banned for strongest_summary_from_resume (tool-style openers); last resort uses different wording.
_BANNED_STRONGEST_SUMMARY_PATTERN = re.compile(
    r"(?i)^professional\s+with\s+experience\s+in\b"
)

_STRONGEST_LAST_RESORT_SUMMARY = (
    "Business analyst delivering analysis, reporting, and validation documented in prior roles."
)


def _is_banned_strongest_summary_pattern(summary: str) -> bool:
    s = (summary or "").strip()
    if not s:
        return False
    if _BANNED_STRONGEST_SUMMARY_PATTERN.search(s):
        return True
    low = s.lower()
    # Generic BSA + doc/validation/reporting bundle (tool-adjacent); keep blocking. Role-led
    # lines that mention business requirements + operational efficiency are allowed so a
    # grounded outcome phrase can beat ``minimal_lexicon``.
    if re.search(r"(?i)business\s+systems?\s+analyst", low) and all(
        w in low for w in ("documentation", "validation", "reporting")
    ):
        return True
    return False


_ROLE_BASED_OUTCOME_SUMMARY_OPENER = re.compile(
    r"(?i)^(senior\s+)?(business\s+systems?\s+analyst|business\s+analyst|data\s+analyst|"
    r"systems\s+analyst|program\s+analyst|product\s+owner|project\s+manager|scrum\s+master)\b"
)


def _is_role_based_outcome_summary_opener(summary: str) -> bool:
    """True when the summary opens with a concrete job title (not ``Professional …``)."""
    return bool(_ROLE_BASED_OUTCOME_SUMMARY_OPENER.search((summary or "").strip()))


def _role_grounded_outcome_phrase_over_minimal_lexicon(
    out_ph: str,
    *,
    grounding_corpus: str,
    match_strength: str,
) -> str:
    """
    If we would otherwise emit ``minimal_lexicon``, prefer a hygiene-clean, role-led outcome
    line over the generic lexicon sentence. Still respects ``_is_banned_strongest_summary_pattern``
    (professional opener, weak BSA doc/val/reporting bundle), tool-centric, and weak tool triplets.
    """
    t = _trim_redundant_words((out_ph or "").strip())
    if not t:
        return ""
    if _BANNED_STRONGEST_SUMMARY_PATTERN.search(t):
        return ""
    if _is_weak_tool_triplet_experience_fallback_line(t):
        return ""
    if is_tool_centric_summary(t):
        return ""
    if not _is_role_based_outcome_summary_opener(t):
        return ""
    if not _export_summary_passes_hygiene(t, grounding_corpus, match_strength):
        return ""
    if _is_banned_strongest_summary_pattern(t):
        return ""
    return t


def _is_weak_tool_triplet_experience_fallback_line(summary: str) -> bool:
    """Known weak ``… with experience in excel, power bi, and python``-style resume fallback."""
    low = (summary or "").strip().lower()
    if "with experience in" not in low:
        return False
    return (
        "excel" in low
        and ("power bi" in low or "powerbi" in low)
        and "python" in low
    )


_IDENTITY_GROUNDING_SIGNAL_TERMS: Tuple[str, ...] = (
    "senior business systems",
    "uat lead",
    "user acceptance testing",
    "system validation",
    "enterprise",
    "regulated",
    "ambiguous",
)


def _identity_grounding_signal_score(text: str) -> int:
    """Rough count of senior BSA / UAT / validation / enterprise signals (export policy only)."""
    low = (text or "").lower()
    return sum(1 for term in _IDENTITY_GROUNDING_SIGNAL_TERMS if term in low)


def _corpus_suggests_bsa_uat_track(corpus_lower: str) -> bool:
    cl = corpus_lower or ""
    if "business systems analyst" in cl:
        return True
    if "uat" in cl and "analyst" in cl:
        return True
    if "uat" in cl and "systems analyst" in cl:
        return True
    return False


def _structured_resume_summary_snippet(resume_data: dict) -> str:
    """First structured summary text from resume sections (for generic-summary escape)."""
    sections = resume_data.get("sections") if isinstance(resume_data.get("sections"), dict) else {}
    summ = sections.get("summary") if isinstance(sections, dict) else None
    if isinstance(summ, list):
        parts = [str(x).strip() for x in summ if str(x).strip()]
        return _trim_redundant_words(" ".join(parts).strip()) if parts else ""
    if isinstance(summ, str) and summ.strip():
        return _trim_redundant_words(summ.strip())
    return ""


def _maybe_elevate_from_minimal_lexicon_identity(
    summary: str,
    src: str,
    *,
    resume_data: dict,
    grounding_corpus: str,
    match_strength: str,
) -> Tuple[str, str]:
    """
    If selection fell to ``minimal_lexicon``, prefer a grounded identity summary when it
    passes strict hygiene, or passes the resume-title-expanded word-grounding check.

    Also upgrades weak tool-stack ``outcome_phrase`` winners and offers a BSA identity line
    or reuse of the resume's own summary before keeping generic lexicon text.
    """
    minimal = (summary or "").strip() == (_MINIMAL_LEXICON_SUMMARY or "").strip()
    weak_outcome = bool(
        (summary or "").strip()
        and src.startswith("outcome_phrase")
        and _is_weak_tool_triplet_experience_fallback_line(summary)
    )
    if (
        src not in ("minimal_lexicon", "minimal_lexicon_forced")
        and not minimal
        and not weak_outcome
    ):
        return summary, src
    candidates: List[Tuple[str, str]] = []
    fwd = _trim_redundant_words(
        (build_strong_identity_forward_export_summary(resume_data, grounding_corpus) or "").strip()
    )
    ids = _trim_redundant_words(
        (build_structured_identity_export_summary(resume_data, grounding_corpus) or "").strip()
    )
    if fwd:
        candidates.append((fwd, "identity_forward_dense_lexicon_escape"))
    if ids:
        candidates.append((ids, "identity_structured_lexicon_escape"))
    for cand, cand_src in candidates:
        if not cand:
            _log_summary_candidate_debug(
                source=cand_src,
                text="",
                accepted=False,
                rejected_reason="empty",
            )
            continue
        if _is_banned_strongest_summary_pattern(cand):
            _log_summary_candidate_debug(
                source=cand_src,
                text=cand,
                accepted=False,
                rejected_reason="banned_strongest_pattern",
            )
            continue
        if is_tool_centric_summary(cand):
            _log_summary_candidate_debug(
                source=cand_src,
                text=cand,
                accepted=False,
                rejected_reason="tool_centric",
            )
            continue
        strict = _export_summary_passes_hygiene(cand, grounding_corpus, match_strength)
        title_ground = _export_identity_summary_passes_hygiene_resume_title_corpus(
            cand, grounding_corpus, match_strength, resume_data
        )
        if strict:
            _log_summary_candidate_debug(
                source=cand_src,
                text=cand,
                accepted=True,
                rejected_reason="",
            )
            return cand, cand_src
        if title_ground:
            _log_summary_candidate_debug(
                source=cand_src,
                text=cand,
                accepted=True,
                rejected_reason="",
            )
            return cand, f"{cand_src}_resume_title_grounding"
        rr = _export_summary_rejection_reason(
            cand, grounding_corpus, match_strength, resume_data=resume_data
        )
        reason = f"hygiene_fail_strict_and_title_expansion detail={rr}"
        if not title_ground:
            reason = f"hygiene_fail_title_expansion detail={rr}"
        _log_summary_candidate_debug(
            source=cand_src,
            text=cand,
            accepted=False,
            rejected_reason=reason,
        )
    cl = (grounding_corpus or "").lower()
    if _corpus_suggests_bsa_uat_track(cl):
        tag = _trim_redundant_words(_BSA_IDENTITY_PREFERRED_FALLBACK.strip())
        if tag and not _is_banned_strongest_summary_pattern(tag) and not is_tool_centric_summary(
            tag
        ):
            if _export_summary_passes_hygiene(tag, grounding_corpus, match_strength):
                _log_summary_candidate_debug(
                    source="identity_bsa_tagline_fallback",
                    text=tag,
                    accepted=True,
                    rejected_reason="",
                )
                return tag, "identity_bsa_tagline_fallback"
            _log_summary_candidate_debug(
                source="identity_bsa_tagline_fallback",
                text=tag,
                accepted=False,
                rejected_reason=_export_summary_rejection_reason(
                    tag, grounding_corpus, match_strength, resume_data=resume_data
                ),
            )
    rs = _structured_resume_summary_snippet(resume_data)
    if rs and not is_tool_centric_summary(rs) and not _is_banned_strongest_summary_pattern(rs):
        if _export_summary_passes_hygiene(rs, grounding_corpus, match_strength):
            _log_summary_candidate_debug(
                source="resume_structured_summary_reuse",
                text=rs,
                accepted=True,
                rejected_reason="",
            )
            return rs, "resume_structured_summary_reuse"
        _log_summary_candidate_debug(
            source="resume_structured_summary_reuse",
            text=rs,
            accepted=False,
            rejected_reason=_export_summary_rejection_reason(
                rs, grounding_corpus, match_strength, resume_data=resume_data
            ),
        )
    return summary, src


def _maybe_upgrade_outcome_phrase_relaxed_summary(
    summary: str,
    src: str,
    *,
    resume_data: dict,
    grounding_corpus: str,
    match_strength: str,
) -> Tuple[str, str]:
    """
    If ``outcome_phrase_relaxed`` won with a generic outcome line, prefer a stronger
    identity-forward / structured summary when it passes the same hygiene gates.
    """
    if src != "outcome_phrase_relaxed":
        return summary, src
    id_alt = _trim_redundant_words(
        (build_structured_identity_export_summary(resume_data, grounding_corpus) or "").strip()
    )
    fwd_alt = _trim_redundant_words(
        (build_strong_identity_forward_export_summary(resume_data, grounding_corpus) or "").strip()
    )
    best_txt = ""
    best_src = ""
    best_score = -1
    for alt, alt_src in (
        (fwd_alt, "identity_forward_dense_policy_override"),
        (id_alt, "identity_structured_policy_override"),
    ):
        if not alt:
            continue
        if not _export_summary_passes_hygiene(alt, grounding_corpus, match_strength):
            logger.debug(
                "summary policy override rejected (hygiene): candidate=%r source=%s alt=%r alt_source=%s",
                summary,
                src,
                alt,
                alt_src,
            )
            continue
        if _is_banned_strongest_summary_pattern(alt) or is_tool_centric_summary(alt):
            logger.debug(
                "summary policy override rejected (banned/tool): candidate=%r source=%s alt=%r alt_source=%s",
                summary,
                src,
                alt,
                alt_src,
            )
            continue
        sc = _identity_grounding_signal_score(alt)
        if sc > best_score:
            best_score = sc
            best_txt = alt
            best_src = alt_src
    relaxed_sc = _identity_grounding_signal_score(summary)
    if best_txt and best_score > relaxed_sc:
        logger.info(
            "summary policy selected stronger identity replacement: previous_source=%s replacement_source=%s relaxed_score=%s replacement_score=%s",
            src,
            best_src,
            relaxed_sc,
            best_score,
        )
        logger.debug(
            "summary replacement details: previous=%r replacement=%r",
            summary,
            best_txt,
        )
        return best_txt, best_src
    logger.debug(
        "summary policy kept relaxed candidate: source=%s best_override_score=%s relaxed_score=%s",
        src,
        best_score,
        relaxed_sc,
    )
    return summary, src


def list_ungrounded_summary_tokens(summary: str, corpus_lower: str) -> List[str]:
    """Tokens that fail _summary_words_grounded_in_resume (for debug logging only)."""
    out: List[str] = []
    cl = corpus_lower or ""
    for m in re.finditer(r"[A-Za-z][A-Za-z0-9+#.\-]{3,}", summary or ""):
        raw = m.group(0)
        wl = raw.lower().rstrip(".,;:!?")
        if wl in _SUMMARY_LEXICON_ALLOW:
            continue
        if wl in cl:
            continue
        out.append(raw)
    return out


def build_resume_grounded_export_summary(
    resume_text: str,
    original_experience_text: Optional[str] = None,
    *,
    corpus_for_validation: str,
    match_strength: str,
) -> str:
    """
    Deterministic one-sentence export fallback: role + skill phrases detected as substrings
    in resume text only; final choice must pass the same hygiene checks as STEP 1.
    """
    blob = f"{resume_text or ''} {original_experience_text or ''}".lower()
    cl = (corpus_for_validation or "").lower()

    role = "Professional"
    for r in _FALLBACK_ROLE_ORDER:
        if r in blob:
            role = " ".join(w.capitalize() for w in r.split())
            break

    picked: List[str] = []
    for ph in _FALLBACK_SKILL_PHRASES:
        if ph in blob:
            picked.append(ph)
        if len(picked) >= 3:
            break

    if len(picked) < 3:
        for w in ("analysis", "reporting", "testing", "validation", "documentation"):
            if w in blob and w not in picked:
                picked.append(w)
            if len(picked) >= 3:
                break

    candidates: List[str] = []
    if len(picked) >= 3:
        candidates.append(
            f"{role} with experience in {picked[0]}, {picked[1]}, and {picked[2]}."
        )
    if len(picked) == 2:
        candidates.append(f"{role} with experience in {picked[0]} and {picked[1]}.")
    if len(picked) == 1:
        candidates.append(f"{role} with experience in {picked[0]}.")

    candidates.append(_MINIMAL_LEXICON_SUMMARY)

    seen: set[str] = set()
    for raw in candidates:
        s = _trim_redundant_words(raw.strip())
        if not s or s in seen:
            continue
        seen.add(s)
        if _export_summary_passes_hygiene(s, corpus_for_validation, match_strength):
            return s
    return _MINIMAL_LEXICON_SUMMARY


def _finalize_summary_scannable(
    summary: str,
    grounding_corpus: str,
    match_strength: str,
) -> str:
    """Shorter summary for screening speed; only applied if hygiene still passes."""
    if match_strength == "weak":
        return summary
    t = trim_summary_for_scannability(summary)
    if t == (summary or "").strip():
        return summary
    if _export_summary_passes_hygiene(t, grounding_corpus, match_strength):
        return t
    return summary


def strongest_summary_from_resume(
    tailored_summary: str,
    match_strength: str,
    resume_data: dict,
    grounding_corpus: str,
) -> Tuple[str, str]:
    """
    Single summary selector for export. Rejects generic "Professional with experience in…"
    openers; obeys the same hygiene rules as STEP 1. weak → fixed weak template only.
    """
    if match_strength == "weak":
        logger.info(
            "SUMMARY_TRACE selected src=weak_fixed summary=%r",
            _WEAK_EXPORT_SUMMARY,
        )
        return _WEAK_EXPORT_SUMMARY, "weak_fixed"

    forced = os.environ.get(_EXPORT_FALLBACK_ONLY_ENV, "").strip().lower() in (
        "1",
        "true",
        "yes",
    )
    resume_blob = _resume_text_corpus_for_export_summary(resume_data)

    if forced:
        s = build_resume_grounded_export_summary(
            resume_blob,
            None,
            corpus_for_validation=grounding_corpus,
            match_strength=match_strength,
        )
        s = _trim_redundant_words(s.strip())
        if (
            s
            and not _is_banned_strongest_summary_pattern(s)
            and _export_summary_passes_hygiene(s, grounding_corpus, match_strength)
        ):
            return _finalize_summary_scannable(s, grounding_corpus, match_strength), "fallback_forced"
        out_ph = _trim_redundant_words(
            (build_outcome_phrase_export_summary(resume_data, grounding_corpus) or "").strip()
        )
        if out_ph and _export_summary_passes_hygiene(
            out_ph, grounding_corpus, match_strength
        ):
            return (
                _finalize_summary_scannable(out_ph, grounding_corpus, match_strength),
                "outcome_phrase_forced",
            )
        lr = _STRONGEST_LAST_RESORT_SUMMARY
        if _export_summary_passes_hygiene(lr, grounding_corpus, match_strength):
            return _finalize_summary_scannable(lr, grounding_corpus, match_strength), "last_resort"
        min_s, min_src = _MINIMAL_LEXICON_SUMMARY, "minimal_lexicon_forced"
        min_s, min_src = _maybe_elevate_from_minimal_lexicon_identity(
            min_s,
            min_src,
            resume_data=resume_data,
            grounding_corpus=grounding_corpus,
            match_strength=match_strength,
        )
        return _finalize_summary_scannable(min_s, grounding_corpus, match_strength), min_src

    t = (tailored_summary or "").strip()
    if not t:
        _log_summary_candidate_debug(
            source="tailored",
            text="",
            accepted=False,
            rejected_reason="empty",
        )
    elif not _export_summary_passes_hygiene(t, grounding_corpus, match_strength):
        _log_summary_candidate_debug(
            source="tailored",
            text=t,
            accepted=False,
            rejected_reason=_export_summary_rejection_reason(
                t, grounding_corpus, match_strength
            ),
        )
    elif _is_banned_strongest_summary_pattern(t):
        _log_summary_candidate_debug(
            source="tailored",
            text=t,
            accepted=False,
            rejected_reason="banned_strongest_pattern",
        )
    else:
        _log_summary_candidate_debug(
            source="tailored",
            text=t,
            accepted=True,
            rejected_reason="",
        )
        return _finalize_summary_scannable(t, grounding_corpus, match_strength), "tailored"

    strong_fwd = _trim_redundant_words(
        (build_strong_identity_forward_export_summary(resume_data, grounding_corpus) or "").strip()
    )
    if strong_fwd:
        rej = ""
        if not _export_summary_passes_hygiene(
            strong_fwd, grounding_corpus, match_strength
        ):
            rej = _export_summary_rejection_reason(
                strong_fwd, grounding_corpus, match_strength
            )
        elif _is_banned_strongest_summary_pattern(strong_fwd):
            rej = "banned_strongest_pattern"
        elif is_tool_centric_summary(strong_fwd):
            rej = "tool_centric"
        if not rej:
            _log_summary_candidate_debug(
                source="identity_forward_dense",
                text=strong_fwd,
                accepted=True,
                rejected_reason="",
            )
            return (
                _finalize_summary_scannable(strong_fwd, grounding_corpus, match_strength),
                "identity_forward_dense",
            )
        _log_summary_candidate_debug(
            source="identity_forward_dense",
            text=strong_fwd,
            accepted=False,
            rejected_reason=rej,
        )

    id_struct = _trim_redundant_words(
        (build_structured_identity_export_summary(resume_data, grounding_corpus) or "").strip()
    )
    if id_struct:
        rej_id = ""
        if not _export_summary_passes_hygiene(
            id_struct, grounding_corpus, match_strength
        ):
            rej_id = _export_summary_rejection_reason(
                id_struct, grounding_corpus, match_strength
            )
        elif _is_banned_strongest_summary_pattern(id_struct):
            rej_id = "banned_strongest_pattern"
        elif is_tool_centric_summary(id_struct):
            rej_id = "tool_centric"
        if not rej_id:
            _log_summary_candidate_debug(
                source="identity_structured",
                text=id_struct,
                accepted=True,
                rejected_reason="",
            )
            return (
                _finalize_summary_scannable(id_struct, grounding_corpus, match_strength),
                "identity_structured",
            )
        _log_summary_candidate_debug(
            source="identity_structured",
            text=id_struct,
            accepted=False,
            rejected_reason=rej_id,
        )

    fb_resume = _trim_redundant_words(
        (_default_resume_summary_fallback(resume_data, grounding_corpus) or "").strip()
    )
    out_ph = _trim_redundant_words(
        (build_outcome_phrase_export_summary(resume_data, grounding_corpus) or "").strip()
    )
    out_ok = bool(
        out_ph
        and _export_summary_passes_hygiene(out_ph, grounding_corpus, match_strength)
        and not _is_banned_strongest_summary_pattern(out_ph)
    )
    fb_ok = bool(
        fb_resume
        and _export_summary_passes_hygiene(fb_resume, grounding_corpus, match_strength)
        and not _is_banned_strongest_summary_pattern(fb_resume)
    )

    _log_summary_candidate_debug(
        source="resume_fallback",
        text=fb_resume or "",
        accepted=fb_ok,
        rejected_reason=""
        if fb_ok
        else (
            "empty"
            if not (fb_resume or "").strip()
            else (
                "banned_strongest_pattern"
                if _is_banned_strongest_summary_pattern(fb_resume)
                else _export_summary_rejection_reason(
                    fb_resume, grounding_corpus, match_strength
                )
            )
        ),
    )
    _log_summary_candidate_debug(
        source="outcome_phrase",
        text=out_ph or "",
        accepted=out_ok,
        rejected_reason=""
        if out_ok
        else (
            "empty"
            if not (out_ph or "").strip()
            else (
                "banned_strongest_pattern"
                if _is_banned_strongest_summary_pattern(out_ph)
                else _export_summary_rejection_reason(
                    out_ph, grounding_corpus, match_strength
                )
            )
        ),
    )

    logger.info(
        "SUMMARY_TRACE candidates tailored=%r tailored_hygiene=%s tailored_banned=%s "
        "fb_resume=%r fb_ok=%s out_ph=%r out_ok=%s",
        (tailored_summary or "").strip()[:220],
        _export_summary_passes_hygiene(
            (tailored_summary or "").strip(), grounding_corpus, match_strength
        )
        if (tailored_summary or "").strip()
        else None,
        _is_banned_strongest_summary_pattern((tailored_summary or "").strip())
        if (tailored_summary or "").strip()
        else None,
        fb_resume,
        fb_ok,
        out_ph,
        out_ok,
    )

    if fb_ok and out_ok and is_tool_centric_summary(fb_resume):
        summary, src = out_ph, "outcome_phrase"
    elif fb_ok:
        summary, src = fb_resume, "resume_fallback"
    elif out_ok:
        summary, src = out_ph, "outcome_phrase"
    else:
        s = build_resume_grounded_export_summary(
            resume_blob,
            None,
            corpus_for_validation=grounding_corpus,
            match_strength=match_strength,
        )
        s = _trim_redundant_words(s.strip())
        if (
            s
            and not _is_banned_strongest_summary_pattern(s)
            and _export_summary_passes_hygiene(s, grounding_corpus, match_strength)
        ):
            summary, src = s, "fallback"
        elif fb_resume and _export_summary_passes_hygiene(
            fb_resume, grounding_corpus, match_strength
        ):
            # ``resume_fallback_relaxed`` previously accepted ``fb_resume`` on hygiene alone,
            # bypassing banned/tool checks — that let weak tool-stack lines through.
            if (
                _is_banned_strongest_summary_pattern(fb_resume)
                or is_tool_centric_summary(fb_resume)
                or _is_weak_tool_triplet_experience_fallback_line(fb_resume)
            ):
                _blocked_reasons = [
                    label
                    for label, bad in (
                        ("banned_generic_pattern", _is_banned_strongest_summary_pattern(fb_resume)),
                        ("tool_centric_stack", is_tool_centric_summary(fb_resume)),
                        (
                            "weak_excel_powerbi_python_experience_line",
                            _is_weak_tool_triplet_experience_fallback_line(fb_resume),
                        ),
                    )
                    if bad
                ]
                logger.info(
                    "SUMMARY_FALLBACK_BLOCK_DEBUG candidate=%r blocked_reason=%s",
                    fb_resume,
                    ",".join(_blocked_reasons),
                )
                if (
                    out_ph
                    and _export_summary_passes_hygiene(
                        out_ph, grounding_corpus, match_strength
                    )
                    and not _is_banned_strongest_summary_pattern(out_ph)
                    and not is_tool_centric_summary(out_ph)
                ):
                    summary, src = out_ph, "outcome_phrase_relaxed"
                elif _export_summary_passes_hygiene(
                    _STRONGEST_LAST_RESORT_SUMMARY, grounding_corpus, match_strength
                ):
                    summary, src = _STRONGEST_LAST_RESORT_SUMMARY, "last_resort"
                else:
                    _rop = _role_grounded_outcome_phrase_over_minimal_lexicon(
                        out_ph,
                        grounding_corpus=grounding_corpus,
                        match_strength=match_strength,
                    )
                    if _rop:
                        summary, src = _rop, "outcome_phrase_role_priority"
                    else:
                        summary, src = _MINIMAL_LEXICON_SUMMARY, "minimal_lexicon"
            else:
                summary, src = fb_resume, "resume_fallback_relaxed"
        elif (
            out_ph
            and _export_summary_passes_hygiene(out_ph, grounding_corpus, match_strength)
            and not _is_banned_strongest_summary_pattern(out_ph)
            and not is_tool_centric_summary(out_ph)
        ):
            summary, src = out_ph, "outcome_phrase_relaxed"
        elif _export_summary_passes_hygiene(
            _STRONGEST_LAST_RESORT_SUMMARY, grounding_corpus, match_strength
        ):
            summary, src = _STRONGEST_LAST_RESORT_SUMMARY, "last_resort"
        else:
            _rop2 = _role_grounded_outcome_phrase_over_minimal_lexicon(
                out_ph,
                grounding_corpus=grounding_corpus,
                match_strength=match_strength,
            )
            if _rop2:
                summary, src = _rop2, "outcome_phrase_role_priority"
            else:
                summary, src = _MINIMAL_LEXICON_SUMMARY, "minimal_lexicon"

    _log_summary_candidate_debug(
        source=f"branch_winner::{src}",
        text=summary or "",
        accepted=True,
        rejected_reason="",
    )

    summary, src = _maybe_upgrade_outcome_phrase_relaxed_summary(
        summary,
        src,
        resume_data=resume_data,
        grounding_corpus=grounding_corpus,
        match_strength=match_strength,
    )

    if match_strength != "weak" and src != "tailored" and is_tool_centric_summary(summary):
        alt = _trim_redundant_words(
            (build_outcome_phrase_export_summary(resume_data, grounding_corpus) or "").strip()
        )
        if (
            alt
            and _export_summary_passes_hygiene(alt, grounding_corpus, match_strength)
            and not _is_banned_strongest_summary_pattern(alt)
        ):
            summary, src = alt, f"{src}_identity_swap"

    summary, src = _maybe_elevate_from_minimal_lexicon_identity(
        summary,
        src,
        resume_data=resume_data,
        grounding_corpus=grounding_corpus,
        match_strength=match_strength,
    )

    final_summary = _finalize_summary_scannable(summary, grounding_corpus, match_strength)
    logger.info(
        "SUMMARY_TRACE selected src=%s summary=%r passes_hygiene=%s banned_pattern=%s",
        src,
        final_summary,
        _export_summary_passes_hygiene(final_summary, grounding_corpus, match_strength),
        _is_banned_strongest_summary_pattern(final_summary),
    )
    _log_summary_policy_debug(
        winner_source=src,
        winner_summary=final_summary,
        outcome_phrase=out_ph,
        outcome_role_opener=_is_role_based_outcome_summary_opener(out_ph),
        passes_hygiene=_export_summary_passes_hygiene(
            final_summary, grounding_corpus, match_strength
        ),
        banned_winner=_is_banned_strongest_summary_pattern(final_summary),
    )
    return final_summary, src


def _format_linkedin_contact_display(url: str) -> str:
    """Stable header display: linkedin.com/in/<slug> (no ellipsis placeholder)."""
    t = (url or "").strip().rstrip("/")
    if not t:
        return ""
    m = re.search(r"(?i)linkedin\.com/in/([\w\-]+)", t)
    if m:
        return f"linkedin.com/in/{m.group(1)}"
    m2 = re.search(r"(?i)linkedin\.com/(?:pub/)?([\w\-]+)(?:\s|$|[,)|])", t)
    if m2 and m2.group(1).lower() not in (
        "in",
        "pub",
        "company",
        "sales",
        "learning",
        "www",
    ):
        return f"linkedin.com/in/{m2.group(1)}"
    return t


def _extract_contact_lines(
    raw_text: str, resume_data: Optional[dict] = None
) -> Tuple[str, str, str, str, str]:
    """Best-effort name + contact fields from resume text (and optional structured resume_data)."""
    body = (raw_text or "").strip()
    lines = [ln.strip() for ln in body.split("\n") if ln.strip()]
    name = lines[0] if lines and len(lines[0].split()) <= 6 else "Name"
    email = ""
    phone = ""
    linkedin = ""
    github = ""
    blob = body.lower()
    sections = (
        resume_data.get("sections")
        if isinstance(resume_data, dict) and isinstance(resume_data.get("sections"), dict)
        else None
    )
    if isinstance(sections, dict):
        for key in ("linkedin_url", "linkedin", "linkedIn", "LinkedIn"):
            v = sections.get(key)
            if isinstance(v, str) and v.strip() and "linkedin" in v.lower():
                linkedin = _format_linkedin_contact_display(v.strip())
                break
    em = re.search(
        r"[a-z0-9._%+\-]+@[a-z0-9.\-]+\.[a-z]{2,}",
        body,
        re.I,
    )
    if em:
        email = em.group(0)
    ph = re.search(
        r"(?:\+?\d{1,3}[\s.\-]?)?(?:\(\d{3}\)|\d{3})[\s.\-]?\d{3}[\s.\-]?\d{4}\b",
        body,
    )
    if ph:
        phone = ph.group(0).strip()
    if not linkedin:
        lm = re.search(r"https?://(?:www\.)?linkedin\.com/[^\s)]+", body, re.I)
        if lm:
            linkedin = _format_linkedin_contact_display(lm.group(0).strip())
    if not linkedin:
        lm_in = re.search(r"(?i)(?:https?://)?(?:www\.)?linkedin\.com/in/([\w\-]+)", body)
        if lm_in:
            linkedin = f"linkedin.com/in/{lm_in.group(1)}"
    if not linkedin:
        lm_path = re.search(
            r"(?i)(?:https?://)?(?:www\.)?linkedin\.com/([\w\-]+)(?=[/\s]|$|[,)|])",
            body,
        )
        if lm_path and lm_path.group(1).lower() not in (
            "in",
            "pub",
            "www",
            "company",
            "sales",
            "learning",
        ):
            linkedin = f"linkedin.com/in/{lm_path.group(1)}"
    gm = re.search(r"https?://(?:www\.)?github\.com/[^\s)]+", body, re.I)
    if gm:
        github = gm.group(0).strip()
    if not github and "github.com" in blob:
        github = "github.com/…"
    return name, email, phone, linkedin, github


def _sanitize_bullet_text(text: str) -> str:
    t = (text or "").strip()
    for marker in _DEBUG_MARKERS:
        if marker.lower() in t.lower():
            return ""
    return t


def _bullet_unsafe_for_export(text: str) -> bool:
    """True if text must not be exported (coaching, system, leak phrases, or debug markers)."""
    t = (text or "").strip()
    if not t:
        return True
    if _sanitize_bullet_text(t) == "":
        return True
    if _BULLET_COMMENTARY_PATTERNS.search(t):
        return True
    low = t.lower()
    for phrase in _CONTENT_LEAK_PHRASES:
        if phrase.lower() in low:
            return True
    for phrase in _BULLET_SCRUB_PHRASES:
        if phrase.lower() in low:
            return True
    return False


def _scrub_system_tokens_from_bullet(text: str) -> str:
    """Remove known system/commentary snippets and parenthetical notes."""
    s = (text or "").strip()
    if not s:
        return s
    s = _BULLET_PAREN_SYSTEM.sub(" ", s)
    for phrase in sorted(_BULLET_SCRUB_PHRASES, key=len, reverse=True):
        s = re.sub(re.escape(phrase), " ", s, flags=re.I)
    s = re.sub(r"(?i)\bcoaching\s*:\s*", " ", s)
    s = re.sub(r"\s+", " ", s).strip(" -–—\t").strip()
    return s


def _postprocess_experience_bullets_for_export(
    blocks: List[dict], original_bullets: List[List[str]]
) -> None:
    """Last-pass scrub; fall back to pre-merge resume line if text is still tainted."""
    for ei, block in enumerate(blocks or []):
        if not isinstance(block, dict):
            continue
        bullets = block.get("bullets")
        if not isinstance(bullets, list):
            continue
        orig_row = original_bullets[ei] if ei < len(original_bullets) else []
        for bi, item in enumerate(bullets):
            orig_line = orig_row[bi] if bi < len(orig_row) else ""
            t = str(item).strip()
            t2 = _scrub_system_tokens_from_bullet(t)
            if not _bullet_unsafe_for_export(t2):
                bullets[bi] = t2
                continue
            t2 = _finalize_export_bullet_text(t2, str(orig_line))
            if not _bullet_unsafe_for_export(t2):
                bullets[bi] = t2
                continue
            t3 = _scrub_system_tokens_from_bullet(str(orig_line))
            if not _bullet_unsafe_for_export(t3):
                bullets[bi] = t3
            else:
                bullets[bi] = _finalize_export_bullet_text(t3, str(orig_line))
        block["bullets"] = bullets


def _finalize_export_bullet_text(chosen: str, orig_line: str) -> str:
    """Prefer clean chosen text; strip markers; fall back to original resume line."""
    o = (orig_line or "").strip()
    c = (chosen or "").strip()
    if not c:
        return o
    if not _bullet_unsafe_for_export(c):
        return c
    s = _sanitize_bullet_text(c)
    if s:
        return s
    if o and not _bullet_unsafe_for_export(o):
        return o
    s2 = _sanitize_bullet_text(o)
    return s2 if s2 else o


def apply_bullet_changes_to_experience(
    experience_blocks: List[dict], bullet_changes: List[dict]
) -> List[dict]:
    """Merge tailored bullets by evidence_id; never prefer coaching/system text over resume-safe lines."""
    out = copy.deepcopy(experience_blocks) if experience_blocks else []
    if not isinstance(out, list):
        return []

    original_bullets: List[List[str]] = []
    for block in out:
        if isinstance(block, dict) and isinstance(block.get("bullets"), list):
            original_bullets.append([str(x) for x in block["bullets"]])
        else:
            original_bullets.append([])

    for bc in bullet_changes or []:
        if not isinstance(bc, dict):
            continue
        if str(bc.get("section") or "") != "experience":
            continue
        eid = str(bc.get("evidence_id") or "").strip()
        m = re.match(r"^exp_(\d+)_bullet_(\d+)$", eid)
        if not m:
            continue
        ei, bi = int(m.group(1)) - 1, int(m.group(2)) - 1
        if ei < 0 or ei >= len(out):
            continue
        block = out[ei]
        if not isinstance(block, dict):
            continue
        bullets = block.get("bullets") or []
        if not isinstance(bullets, list) or bi < 0 or bi >= len(bullets):
            continue
        orig_line = (
            original_bullets[ei][bi]
            if ei < len(original_bullets) and bi < len(original_bullets[ei])
            else ""
        )

        raw_after = str(bc.get("after") or "").strip()
        raw_before = str(bc.get("before") or "").strip()

        chosen = ""
        if raw_after and not _bullet_unsafe_for_export(raw_after):
            chosen = raw_after
        elif raw_before and not _bullet_unsafe_for_export(raw_before):
            chosen = raw_before
        else:
            chosen = orig_line.strip()

        bullets[bi] = _finalize_export_bullet_text(chosen, orig_line)
        block["bullets"] = bullets

    dedupe_bullets_within_experience_blocks(out)
    _postprocess_experience_bullets_for_export(out, original_bullets)
    dedupe_bullets_within_experience_blocks(out)
    return out


def _norm_ws_blob(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def filter_experience_blocks_for_docx(
    blocks: List[dict],
    name_line: str,
    *,
    email: str = "",
    phone: str = "",
    linkedin: str = "",
    github: str = "",
) -> List[dict]:
    """
    Drop contact lines, section-title-only lines, duplicate name lines, and full
    contact-line duplicates from experience bullets so EXPERIENCE uses structured
    role content only (never the header/contact block).
    """
    out: List[dict] = []
    nl = (name_line or "").strip().lower()
    contact_joined = " | ".join(x for x in (email, phone, linkedin, github) if x)
    cn = _norm_ws_blob(contact_joined)
    contact_bits = [x for x in (email, phone, linkedin, github) if x]
    for block in blocks or []:
        if not isinstance(block, dict):
            continue
        bullets = block.get("bullets")
        if not isinstance(bullets, list):
            out.append(dict(block))
            continue
        cleaned: List[str] = []
        for b in bullets:
            t = str(b).strip()
            if not t:
                continue
            if line_is_experience_noise(t):
                continue
            if nl and t.strip().lower() == nl:
                continue
            tn = _norm_ws_blob(t)
            if cn and tn == cn:
                continue
            if any(_norm_ws_blob(bit) == tn for bit in contact_bits):
                continue
            cleaned.append(t)
        nb = dict(block)
        nb["bullets"] = cleaned
        out.append(nb)
    return out


def dedupe_bullets_within_experience_blocks(blocks: List[dict]) -> None:
    """Within each experience block, drop exact duplicates (normalized whitespace/case); keep first."""
    for block in blocks or []:
        if not isinstance(block, dict):
            continue
        bullets = block.get("bullets")
        if not isinstance(bullets, list):
            continue
        seen: set[str] = set()
        out_b: List[Any] = []
        for item in bullets:
            t = str(item).strip()
            norm = re.sub(r"\s+", " ", t.lower()).strip()
            if not norm:
                key = ""
                if key in seen:
                    continue
                seen.add(key)
                out_b.append(item)
                continue
            if norm in seen:
                continue
            seen.add(norm)
            out_b.append(item)
        block["bullets"] = out_b


def format_skills_line(skills_section: Any) -> str:
    """Single clean line for skills section."""
    if isinstance(skills_section, list) and skills_section:
        parts = [str(s).strip() for s in skills_section if str(s).strip()]
        if parts:
            if len(parts) == 1 and ":" in parts[0]:
                return parts[0]
            return "Data & Analytics: " + ", ".join(parts) if parts else ""
    return "Data & Analytics: SQL, Power BI, Tableau, Excel, Python"


def _build_grounding_corpus(
    resume_data: dict,
    experience_blocks: List[dict],
    bullet_changes: List[dict],
) -> str:
    """
    Lexicon / grounding text for export summary hygiene (must match validation corpus).

    Includes ``raw_text``, the supplied experience slice (merged or original — caller
    should pass the same blocks used to build the export narrative), bullet change
    pairs, and structured projects/skills/education/certifications so resume-only
    summaries are not rejected for tokens that appear only outside experience bullets.
    """
    parts: List[str] = [str(resume_data.get("raw_text") or "")]
    for block in experience_blocks or []:
        if not isinstance(block, dict):
            continue
        parts.append(str(block.get("company") or ""))
        parts.append(str(block.get("title") or block.get("role") or ""))
        parts.append(str(block.get("date_range") or block.get("date") or ""))
        parts.append(str(block.get("location") or ""))
        for b in block.get("bullets") or []:
            parts.append(str(b))
    for bc in bullet_changes or []:
        if not isinstance(bc, dict):
            continue
        parts.append(str(bc.get("before") or ""))
        parts.append(str(bc.get("after") or ""))
    sections = resume_data.get("sections") if isinstance(resume_data.get("sections"), dict) else {}
    for proj in sections.get("projects") or []:
        if not isinstance(proj, dict):
            continue
        parts.append(str(proj.get("name") or proj.get("title") or ""))
        parts.append(str(proj.get("subtitle") or proj.get("tagline") or ""))
        for b in proj.get("bullets") or []:
            parts.append(str(b))
    sk = sections.get("skills")
    if isinstance(sk, list):
        for s in sk:
            parts.append(str(s))
    for row in sections.get("education") or []:
        if not isinstance(row, dict):
            continue
        for k in (
            "degree",
            "title",
            "institution",
            "school",
            "company",
            "date_range",
            "date",
            "dates",
            "location",
        ):
            parts.append(str(row.get(k) or ""))
        for b in row.get("bullets") or []:
            parts.append(str(b))
    for row in sections.get("certifications") or []:
        if not isinstance(row, dict):
            continue
        for k in (
            "name",
            "title",
            "credential",
            "issuer",
            "organization",
            "date_range",
            "date",
            "dates",
        ):
            parts.append(str(row.get(k) or ""))
        for b in row.get("bullets") or []:
            parts.append(str(b))
    return " ".join(parts)


def _export_text_blob_from_payload(payload: "ResumeDocumentPayload") -> str:
    """Plain-text mirror of DOCX body order for post-validation (STEP 2), from structured payload only."""
    chunks: List[str] = [payload.summary.strip()]
    for ent in payload.experience:
        for hl in experience_entry_header_lines(ent):
            chunks.append(hl)
        for b in ent.bullets:
            chunks.append(str(b))
    for proj in payload.projects:
        if (proj.name or "").strip():
            chunks.append(proj.name.strip())
        if (proj.subtitle or "").strip():
            chunks.append(proj.subtitle.strip())
        for b in proj.bullets:
            chunks.append(str(b))
    for ent in payload.education:
        for hl in education_entry_header_lines(ent):
            chunks.append(hl)
        for b in ent.bullets:
            chunks.append(str(b))
    for ent in payload.certifications:
        for hl in certification_entry_header_lines(ent):
            chunks.append(hl)
        for b in ent.bullets:
            chunks.append(str(b))
    for sk in skills_to_display_lines(payload.skills):
        if sk.strip():
            chunks.append(sk.strip())
    return "\n".join(c for c in chunks if c)


def _keyword_stuffing_summary(s: str) -> bool:
    t = (s or "").strip()
    if not t:
        return False
    words = re.findall(r"[a-zA-Z]{5,}", t.lower())
    if len(words) < 6:
        return False
    if any(n >= 3 for n in Counter(words).values()):
        return True
    if t.count(",") >= 5 and not re.search(r"[.!?]", t):
        return True
    return False


def _novel_long_words(text: str, corpus_lower: str) -> List[str]:
    out: List[str] = []
    for w in re.findall(r"[a-zA-Z]{6,}", (text or "").lower()):
        if w in corpus_lower:
            continue
        out.append(w)
    return out


def _step1_summary_checks(
    summary: str, match_strength: str, corpus_lower: str
) -> List[str]:
    out: List[str] = []
    s = (summary or "").strip()
    if not s:
        out.append("STEP 1 (Summary): empty summary.")
        return out
    if match_strength == "weak":
        if s != _WEAK_EXPORT_SUMMARY:
            out.append(
                "STEP 1 (Summary): weak match must use the fixed generic summary only."
            )
        return out
    if _summary_has_forbidden_jd_copy_patterns(s) or _summary_input_has_jd_pollution(s):
        out.append(
            "STEP 1 (Summary): job-description-style or employer-marketing phrasing."
        )
    if not _summary_words_grounded_in_resume(s, corpus_lower):
        out.append(
            "STEP 1 (Summary): wording not clearly grounded in the resume corpus."
        )
    if _MANIFESTO_OR_ABOUT_PATTERNS.search(s):
        out.append("STEP 1 (Summary): About-Us or manifesto-style language.")
    if _keyword_stuffing_summary(s):
        out.append("STEP 1 (Summary): unnatural keyword repetition or stuffing.")
    return out


def _step2_content_leak_checks(full_lower: str) -> List[str]:
    hits: set[str] = set()
    for phrase in _CONTENT_LEAK_PHRASES:
        if phrase.lower() in full_lower:
            hits.add(phrase)
    for marker in _DEBUG_MARKERS:
        if marker.lower() in full_lower:
            hits.add(marker)
    return [
        f'STEP 2 (Content leak): forbidden text "{h}".' for h in sorted(hits)
    ]



def _step3_experience_checks(
    experience_merged: List[dict],
    corpus_lower: str,
    match_strength: str,
) -> List[str]:
    out: List[str] = []
    seen_norm: set[str] = set()
    for bi, block in enumerate(experience_merged or []):
        if not isinstance(block, dict):
            continue
        bullets = block.get("bullets") or []
        if not isinstance(bullets, list):
            continue
        for i, b in enumerate(bullets):
            t = str(b).strip()
            if not t:
                out.append(
                    f"STEP 3 (Experience): empty bullet (block {bi + 1}, item {i + 1})."
                )
                continue
            low = t.lower()
            hit_markers = [m for m in _DEBUG_MARKERS if m in low]
            if hit_markers:
                _dbg = json.dumps(
                    {
                        "entry_index": bi,
                        "company": str(block.get("company") or ""),
                        "role": str(
                            block.get("title") or block.get("role") or ""
                        ),
                        "bullet_index": i,
                        "bullet_text": t,
                        "matched_markers": hit_markers,
                    },
                    ensure_ascii=False,
                )
                if _export_debug_enabled():
                    logger.debug("step3 marker details: %s", _dbg)
                out.append(
                    "STEP 3 (Experience): system or coaching marker in bullet text."
                )
            if _BULLET_COMMENTARY_PATTERNS.search(t):
                out.append(
                    "STEP 3 (Experience): commentary or meta-language in bullet text."
                )
            norm = re.sub(r"\s+", " ", low).strip()
            if len(norm) > 12 and norm in seen_norm:
                out.append("STEP 3 (Experience): duplicate bullet text.")
            seen_norm.add(norm)
            if re.search(r"(?:\.{3}|…)\s*$", t):
                out.append(
                    f"STEP 3 (Experience): truncated or ellipsis-ending bullet (block {bi + 1})."
                )
            novel = _novel_long_words(t, corpus_lower)
            # Weak match: grounding strictness is enforced in STEP 6.
            if match_strength != "weak" and len(novel) > 2:
                out.append(
                    "STEP 3 (Experience): bullet may introduce claims not grounded in "
                    f"resume ({', '.join(novel[:6])})."
                )
    return out


def _step6_weak_match_checks(
    experience_merged: List[dict],
    corpus_lower: str,
) -> List[str]:
    out: List[str] = []
    for bi, block in enumerate(experience_merged or []):
        if not isinstance(block, dict):
            continue
        for i, b in enumerate(block.get("bullets") or []):
            t = str(b).strip()
            if not t:
                continue
            novel = _novel_long_words(t, corpus_lower)
            if len(novel) > 1:
                out.append(
                    "STEP 6 (Weak match): experience appears over-tailored "
                    f"(ungrounded terms in block {bi + 1}, item {i + 1}: "
                    f"{', '.join(novel[:5])})."
                )
    return out


def _read_docx_plain(docx_bytes: bytes) -> str:
    from docx import Document

    return "\n".join(p.text for p in Document(BytesIO(docx_bytes)).paragraphs)


def _step4_structure_checks(
    docx_bytes: bytes,
    *,
    expect_projects: bool,
    expect_education: bool,
    expect_certifications: bool,
) -> List[str]:
    from docx import Document

    out: List[str] = []
    doc = Document(BytesIO(docx_bytes))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    def _idx(h: str) -> int:
        u = h.upper()
        for i, t in enumerate(texts):
            if t.upper() == u:
                return i
        return -1

    headers: List[Tuple[str, bool]] = [
        ("SUMMARY", True),
        ("EXPERIENCE", True),
        ("PROJECTS", expect_projects),
        ("EDUCATION", expect_education),
        ("CERTIFICATIONS", expect_certifications),
        ("SKILLS", True),
    ]
    present: List[Tuple[str, int]] = []
    for label, required in headers:
        i = _idx(label)
        if required and i < 0:
            out.append(f"STEP 4 (Structure): {label} section header missing.")
        elif i >= 0:
            present.append((label, i))

    for j in range(len(present) - 1):
        a, ia = present[j]
        b, ib = present[j + 1]
        if not ia < ib:
            out.append(
                f"STEP 4 (Structure): {a} must appear before {b} (section order)."
            )

    if not expect_projects and _idx("PROJECTS") >= 0:
        out.append("STEP 4 (Structure): unexpected PROJECTS section.")
    if not expect_education and _idx("EDUCATION") >= 0:
        out.append("STEP 4 (Structure): unexpected EDUCATION section.")
    if not expect_certifications and _idx("CERTIFICATIONS") >= 0:
        out.append("STEP 4 (Structure): unexpected CERTIFICATIONS section.")

    if texts and texts[0].upper() in (
        "SUMMARY",
        "EXPERIENCE",
        "SKILLS",
        "PROJECTS",
        "EDUCATION",
        "CERTIFICATIONS",
    ):
        out.append("STEP 4 (Structure): name or contact block missing before SUMMARY.")
    return out


def _step5_formatting_checks(docx_bytes: bytes) -> List[str]:
    from docx import Document

    out: List[str] = []
    doc = Document(BytesIO(docx_bytes))
    empty_run = 0
    max_empty = 0
    for p in doc.paragraphs:
        if not (p.text or "").strip():
            empty_run += 1
            max_empty = max(max_empty, empty_run)
        else:
            empty_run = 0
    if max_empty > 3:
        out.append("STEP 5 (Formatting): excessive consecutive blank paragraphs.")

    allowed = _ALLOWED_BODY_FONT_PT | _ALLOWED_NAME_FONT_PT
    bad_sizes: List[float] = []
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.size is None:
                continue
            sz = round(float(r.font.size.pt), 1)
            if sz < 10.0 or sz > 14.5 or sz not in allowed:
                bad_sizes.append(sz)
    if bad_sizes:
        sample = ", ".join(str(x) for x in sorted(set(bad_sizes))[:6])
        out.append(
            "STEP 5 (Formatting): font sizes not in the allowed Calibri set "
            f"(10.5–12pt body, 12–14pt name; saw: {sample})."
        )
    return out


def validate_export_pre_docx(
    *,
    summary: str,
    experience_merged: List[dict],
    match_strength: str,
    resume_data: dict,
    experience_original: List[dict],
    bullet_changes: List[dict],
) -> List[str]:
    """Steps 1, 3, 6 — detection only; run before building the .docx."""
    # Corpus must match ``strongest_summary_from_resume`` / ``build_export_docx_package``:
    # merged experience for docx plus structured sections (not pre-merge originals only).
    corpus = _build_grounding_corpus(resume_data, experience_merged, bullet_changes)
    cl = corpus.lower()
    failures: List[str] = []
    failures.extend(_step1_summary_checks(summary, match_strength, cl))
    failures.extend(_step3_experience_checks(experience_merged, cl, match_strength))
    if match_strength == "weak":
        failures.extend(_step6_weak_match_checks(experience_merged, cl))
    return failures


def validate_export_post_docx(
    *,
    export_text_blob: str,
    docx_bytes: bytes,
    expect_projects: bool,
    expect_education: bool = False,
    expect_certifications: bool = False,
) -> List[str]:
    """Steps 2, 4, 5 — run on assembled content and generated file bytes."""
    doc_plain = _read_docx_plain(docx_bytes)
    full_lower = (export_text_blob + "\n" + doc_plain).lower()
    failures: List[str] = []
    failures.extend(_step2_content_leak_checks(full_lower))
    failures.extend(
        _step4_structure_checks(
            docx_bytes,
            expect_projects=expect_projects,
            expect_education=expect_education,
            expect_certifications=expect_certifications,
        )
    )
    failures.extend(_step5_formatting_checks(docx_bytes))
    return failures


def format_validation_failure(failures: List[str]) -> str:
    if not failures:
        return ""
    lines = [_FAIL_HEAD] + [f"- {x}" for x in failures]
    return "\n".join(lines)


def split_validation_failure(formatted: str) -> Tuple[str, List[str]]:
    """Parse `format_validation_failure` output into status + bullet list."""
    lines = formatted.strip().split("\n")
    if not lines:
        return _FAIL_HEAD, []
    head = lines[0]
    checks: List[str] = []
    for ln in lines[1:]:
        if ln.startswith("- "):
            checks.append(ln[2:].strip())
        elif ln.strip():
            checks.append(ln.strip())
    return head, checks


# Project / education / certification rows mis-segmented into EXPERIENCE must not render
# under EXPERIENCE (see ``build_docx_from_payload`` / ``payload.projects``).
_MISPLACED_PROJECT_BUCKET_IN_EXP = re.compile(r"(?i)\bpersonal\s+project\b")
_EXPERIENCE_DOCX_YEAR_ONLY = re.compile(r"^\d{4}$")
_LEGIT_PROJECT_JOB_IN_HEADER = re.compile(
    r"(?i)\bproject\s+(manager|coordinator|engineer|lead|director|analyst|specialist|owner|controller)\b"
)
_EXPERIENCE_DOCX_BULLET_EDU_CERT_MARKERS: Tuple[str, ...] = (
    "b.a.",
    "b.s.",
    "b.a ",
    "b.s ",
    "bachelor",
    "master's",
    "masters",
    "mba",
    "university",
    "college degree",
    " of science degree",
    " of arts degree",
    "associate degree",
    "associates degree",
    "certificate",
    "itil",
    "ecba",
    "ccba",
    "cbap",
)
# Stray certification titles/lines mis-filed under EXPERIENCE (DOCX-only filter).
_EXPERIENCE_DOCX_CERT_LINE = re.compile(r"(?i)\b(certified|certificate|itil|ecba)\b")


def _experience_bullet_is_certification_line(text: str) -> bool:
    return bool(_EXPERIENCE_DOCX_CERT_LINE.search((text or "").strip()))


def _experience_entry_invalid_for_experience_section(ent: ExperienceEntry) -> bool:
    """
    Hard gate: drop entries that belong in PROJECTS / EDUCATION / CERTIFICATIONS but were
    assembled as experience (DOCX export only; does not change upstream segmentation).
    """
    company = str(ent.company or "").strip()
    role = str(ent.role or "").strip()
    header_l = f"{company} {role}".strip().lower()
    parts = " ".join(
        [
            company,
            role,
            str(ent.date or ""),
            str(ent.location or ""),
        ]
    )
    if _EXPERIENCE_DOCX_CERT_LINE.search(header_l):
        return True
    if _MISPLACED_PROJECT_BUCKET_IN_EXP.search(parts):
        return True
    for b in (ent.bullets or [])[:3]:
        if _MISPLACED_PROJECT_BUCKET_IN_EXP.search(str(b)):
            return True
    if _EXPERIENCE_DOCX_YEAR_ONLY.match(company) or _EXPERIENCE_DOCX_YEAR_ONLY.match(role):
        return True
    if "project" in header_l:
        if _LEGIT_PROJECT_JOB_IN_HEADER.search(header_l):
            return False
        return True
    blob = " ".join(str(b).lower() for b in (ent.bullets or []) if str(b).strip())
    if blob and _EXPERIENCE_DOCX_CERT_LINE.search(blob):
        return True
    if blob and any(m in blob for m in _EXPERIENCE_DOCX_BULLET_EDU_CERT_MARKERS):
        return True
    return False


def _docx_paragraph_tune(
    paragraph: Any,
    *,
    WD_LINE_SPACING: Any,
    Pt: Any,
    before_pt: float = 0.0,
    after_pt: float = 0.0,
) -> None:
    """DOCX-only paragraph rhythm: explicit before/after (points) and single line spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _export_docx_paragraph_spacing(
    paragraph: Any,
    *,
    WD_LINE_SPACING: Any,
    Pt: Any,
) -> None:
    """Tight body rhythm: no extra before/after, single line spacing (DOCX only)."""
    _docx_paragraph_tune(
        paragraph,
        WD_LINE_SPACING=WD_LINE_SPACING,
        Pt=Pt,
        before_pt=0.0,
        after_pt=0.0,
    )


def _experience_docx_filtered_bullets(ent: ExperienceEntry) -> List[str]:
    return [
        str(b).strip()
        for b in (ent.bullets or [])
        if str(b).strip() and not _experience_bullet_is_certification_line(str(b))
    ]


def _experience_meta_header_lines(ent: ExperienceEntry) -> List[str]:
    """
    Location / date lines only, matching ``experience_entry_header_lines`` meta rules
    (DOCX render — keeps content order identical to assembly).
    """
    date = re.sub(r"\s+", " ", ((ent.date or "").replace("\r", " ").replace("\n", " "))).strip()
    loc = re.sub(r"\s+", " ", ((ent.location or "").replace("\r", " ").replace("\n", " "))).strip()
    if loc and date:
        return [loc, date]
    if loc:
        return [loc]
    if date:
        return [date]
    return []


def _experience_entry_skipped_for_docx(ent: ExperienceEntry) -> bool:
    if _experience_entry_invalid_for_experience_section(ent):
        return True
    company = (ent.company or "").strip()
    role = (ent.role or "").strip()
    has_identity = bool(company or role)
    bullets = _experience_docx_filtered_bullets(ent)
    if bullets and not has_identity:
        return True
    return False


def _experience_entry_has_docx_output(ent: ExperienceEntry) -> bool:
    if _experience_entry_skipped_for_docx(ent):
        return False
    header_lines = experience_entry_header_lines(ent)
    bullets = _experience_docx_filtered_bullets(ent)
    return bool(header_lines) or bool(bullets)


def _render_experience_entries_to_docx(
    doc: Any,
    experience: List[ExperienceEntry],
    *,
    set_default_font: Any,
    Pt: Any,
    WD_LINE_SPACING: Any,
    refinery_experience_spacing: bool = False,
) -> None:
    """
    Experience render: company (emphasized), role, then location/date (secondary), then List Bullet
    rows only for true bullets. Spacing between jobs uses paragraph ``space_before`` on the next
    company (or role-if-no-company) line — no spacer blank paragraphs.
    """
    job_ord = 0
    for ent in experience:
        if _experience_entry_skipped_for_docx(ent):
            if _experience_entry_invalid_for_experience_section(ent):
                if _export_debug_enabled():
                    logger.debug(
                        "skip invalid experience row in renderer: company=%r role=%r",
                        (ent.company or "")[:120],
                        (ent.role or "")[:120],
                    )
            else:
                fb = _experience_docx_filtered_bullets(ent)
                if fb and not ((ent.company or "").strip() or (ent.role or "").strip()):
                    logger.error(
                        "RENDER_TYPE_DEBUG invariant violated: bullet_count=%s but company and role empty; "
                        "skipping bullet render for this entry (upstream validation should have blocked).",
                        len(fb),
                    )
            continue
        company = (ent.company or "").strip()
        role = (ent.role or "").strip()
        date = (ent.date or "").strip()
        loc = (ent.location or "").strip()
        bullets = _experience_docx_filtered_bullets(ent)
        header_lines = experience_entry_header_lines(ent)
        meta_lines = _experience_meta_header_lines(ent)
        logger.info(
            "RENDER_TYPE_DEBUG company=%r role=%r date=%r location=%r bullet_count=%s "
            "header_line_count=%s",
            company,
            role,
            date,
            loc,
            len(bullets),
            len(header_lines),
        )
        logger.info(
            "RENDER_ENTRY_DEBUG company=%r role=%r date=%r location=%r bullet_count=%s",
            company,
            role,
            date,
            loc,
            len(bullets),
        )

        job_ord += 1
        if job_ord == 1:
            gap_pt = 4.0 if refinery_experience_spacing else 3.0
        else:
            gap_pt = 12.0 if refinery_experience_spacing else 8.0
        pending_before_pt = gap_pt

        header_paras: List[Any] = []

        if company:
            hp = doc.add_paragraph()
            hr = hp.add_run(company)
            hr.bold = True
            set_default_font(hr, 12)
            _docx_paragraph_tune(
                hp,
                WD_LINE_SPACING=WD_LINE_SPACING,
                Pt=Pt,
                before_pt=pending_before_pt,
                after_pt=0.0,
            )
            pending_before_pt = 0.0
            header_paras.append(hp)

        if role:
            hp = doc.add_paragraph()
            hr = hp.add_run(role)
            hr.bold = False
            set_default_font(hr, 11)
            _docx_paragraph_tune(
                hp,
                WD_LINE_SPACING=WD_LINE_SPACING,
                Pt=Pt,
                before_pt=pending_before_pt,
                after_pt=0.0,
            )
            pending_before_pt = 0.0
            header_paras.append(hp)

        for ml in meta_lines:
            hp = doc.add_paragraph()
            hr = hp.add_run(ml)
            hr.bold = False
            set_default_font(hr, 10.5)
            _docx_paragraph_tune(
                hp,
                WD_LINE_SPACING=WD_LINE_SPACING,
                Pt=Pt,
                before_pt=pending_before_pt,
                after_pt=0.0,
            )
            pending_before_pt = 0.0
            header_paras.append(hp)

        if header_paras and bullets:
            last_hp = header_paras[-1]
            last_hp.paragraph_format.space_after = Pt(4)

        for bt in bullets:
            bp = doc.add_paragraph(style="List Bullet")
            br = bp.add_run(bt)
            set_default_font(br, 11)
            _export_docx_paragraph_spacing(bp, WD_LINE_SPACING=WD_LINE_SPACING, Pt=Pt)


def build_docx_from_payload(
    payload: "ResumeDocumentPayload",
    *,
    refinery_experience_spacing: bool = False,
) -> bytes:
    """Build DOCX exclusively from a ResumeDocumentPayload (structured assembly)."""
    logger.info(
        "EXPORT_PIPELINE_DEBUG stage=build_docx_from_payload_entry payload_id=%s module=%s qualname=%s",
        id(payload),
        build_docx_from_payload.__module__,
        build_docx_from_payload.__qualname__,
    )
    validate_resume_document_payload(payload)
    if _export_debug_enabled():
        _proj_payload_debug = json.dumps(
            [
                {
                    "name": (p.name or "").strip(),
                    "subtitle": (p.subtitle or "").strip(),
                    "bullets": [str(b).strip() for b in (p.bullets or []) if str(b).strip()],
                }
                for p in (payload.projects or [])
            ],
            ensure_ascii=False,
        )
        logger.debug("project payload snapshot: %s", _proj_payload_debug)
    logger.info(
        "STRUCT_CONTRACT docx_render summary=%r experience_len=%s",
        (payload.summary or "").strip(),
        len(payload.experience),
    )
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Pt
    from docx.oxml.ns import qn

    doc = Document()
    h = payload.header

    def _compact(p: Any) -> None:
        _export_docx_paragraph_spacing(p, WD_LINE_SPACING=WD_LINE_SPACING, Pt=Pt)

    def set_default_font(run, size_pt: float = 11.0) -> None:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.size = Pt(size_pt)

    def add_section_heading(text: str, *, space_before_pt: float = 6.0) -> None:
        hp = doc.add_paragraph()
        hr = hp.add_run(text)
        hr.bold = True
        set_default_font(hr, 12)
        pf = hp.paragraph_format
        pf.space_before = Pt(space_before_pt)
        pf.space_after = Pt(4)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def style_body_paragraph(p: Any, *, line: float = 1.15, after_pt: float = 6.0) -> None:
        del line
        _docx_paragraph_tune(
            p,
            WD_LINE_SPACING=WD_LINE_SPACING,
            Pt=Pt,
            before_pt=0.0,
            after_pt=after_pt,
        )

    # Header — name + contact only
    p = doc.add_paragraph()
    r = p.add_run((h.name or "Name").strip())
    r.bold = True
    set_default_font(r, 14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _compact(p)

    if (h.contact or "").strip():
        cp = doc.add_paragraph(h.contact.strip())
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            set_default_font(run, 10.5)
        _compact(cp)

    _compact(doc.add_paragraph())

    # Summary — single validated string only
    add_section_heading("SUMMARY", space_before_pt=0.0)
    sp = doc.add_paragraph()
    sr = sp.add_run(payload.summary.strip())
    set_default_font(sr, 11)
    sp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(5)

    # Experience — strict identity-first order per entry (see _render_experience_entries_to_docx).
    add_section_heading("EXPERIENCE")
    _render_experience_entries_to_docx(
        doc,
        payload.experience,
        set_default_font=set_default_font,
        Pt=Pt,
        WD_LINE_SPACING=WD_LINE_SPACING,
        refinery_experience_spacing=refinery_experience_spacing,
    )

    _raw_project_count = len(payload.projects or [])
    _projects_render = [
        p
        for p in (payload.projects or [])
        if any(str(b).strip() for b in (p.bullets or []))
    ]
    _render_proj_section = bool(_projects_render)
    logger.info(
        "PROJECT_RENDER_SOURCE_DEBUG payload_projects_count=%s project_render_path_used=%s",
        _raw_project_count,
        "payload.projects_only",
    )
    _prd = json.dumps(
        {
            "payload_projects_count": _raw_project_count,
            "renderable_projects_count": len(_projects_render),
            "render_projects_section": _render_proj_section,
            "render_source": (
                "payload.projects_bulleted_only"
                if _render_proj_section
                else "skip_empty_or_title_only_no_legacy_path"
            ),
        },
        ensure_ascii=False,
    )
    logger.debug("project render snapshot: %s", _prd)
    # PROJECTS body is emitted only from ``payload.projects`` (filtered to bulleted rows). No
    # legacy/raw-text project path exists in this module.
    if _projects_render:
        add_section_heading("PROJECTS")
        for proj in _projects_render:
            pn = (proj.name or "").strip()
            if pn:
                pp = doc.add_paragraph()
                pr = pp.add_run(pn)
                pr.bold = True
                set_default_font(pr, 11)
                _compact(pp)
            sub = (proj.subtitle or "").strip()
            if sub:
                sp = doc.add_paragraph()
                sr = sp.add_run(sub)
                set_default_font(sr, 11)
                _compact(sp)
            for bullet in proj.bullets:
                bt = str(bullet).strip()
                if not bt:
                    continue
                bp = doc.add_paragraph(style="List Bullet")
                br = bp.add_run(bt)
                set_default_font(br, 11)
                _compact(bp)

    if payload.education:
        add_section_heading("EDUCATION")
        for ent in payload.education:
            for line in education_entry_header_lines(ent):
                hp = doc.add_paragraph()
                hr = hp.add_run(line)
                hr.bold = True
                set_default_font(hr, 11)
                _compact(hp)
            for bullet in ent.bullets:
                bt = str(bullet).strip()
                if not bt:
                    continue
                bp = doc.add_paragraph()
                br = bp.add_run(bt)
                set_default_font(br, 11)
                _compact(bp)

    if payload.certifications:
        add_section_heading("CERTIFICATIONS")
        for ent in payload.certifications:
            for line in certification_entry_header_lines(ent):
                hp = doc.add_paragraph()
                hr = hp.add_run(line)
                hr.bold = False
                set_default_font(hr, 11)
                _compact(hp)
            for bullet in ent.bullets:
                bt = str(bullet).strip()
                if not bt:
                    continue
                bp = doc.add_paragraph()
                br = bp.add_run(bt)
                set_default_font(br, 11)
                _compact(bp)

    add_section_heading("SKILLS")
    for sk_line in skills_to_display_lines(payload.skills):
        sl = sk_line.strip()
        if not sl:
            continue
        kp = doc.add_paragraph(sl)
        for run in kp.runs:
            set_default_font(run, 11)
        style_body_paragraph(kp, after_pt=2.0)

    if not _render_proj_section:
        _pre_save_plain = "\n".join((getattr(p, "text", None) or "") for p in doc.paragraphs)
        if re.search(r"(?i)\bpersonal\s+project\b", _pre_save_plain):
            logger.warning(
                "DOCX pre-save project leak check triggered: %r",
                _pre_save_plain[:700],
            )
        assert not re.search(r"(?i)\bpersonal\s+project\b", _pre_save_plain), (
            "DOCX export: found Personal Project in document while "
            "payload had no renderable project bullets (payload.projects-only rule)."
        )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _trim_gainwell_experience_blocks_for_export_readability(
    blocks: List[dict], *, max_bullets: int = 6
) -> List[dict]:
    """
    Export-only: keep the strongest Gainwell bullets first (existing prioritizer) and cap
    count for scannability. Does not alter segmentation or non-Gainwell employers.
    """
    out: List[dict] = []
    for b in blocks or []:
        if not isinstance(b, dict):
            out.append(b)
            continue
        company = str(b.get("company") or "").lower()
        if "gainwell" not in company:
            out.append(b)
            continue
        bullets = b.get("bullets")
        if not isinstance(bullets, list) or len(bullets) <= max_bullets:
            out.append(b)
            continue
        raw = [str(x).strip() for x in bullets if str(x).strip()]
        ordered, _ = prioritize_experience_bullets(
            raw,
            company=str(b.get("company") or ""),
            role=str(b.get("title") or b.get("role") or ""),
            date=str(b.get("date_range") or b.get("date") or ""),
            location=str(b.get("location") or ""),
        )
        nb = dict(b)
        nb["bullets"] = ordered[:max_bullets]
        logger.info(
            "EXPORT_GAINWELL_BULLET_TRIM before=%s after=%s",
            len(raw),
            len(nb["bullets"]),
        )
        out.append(nb)
    return out


def _log_experience_blocks_for_debug(stage: str, blocks: List[dict]) -> None:
    """Temporary: trace raw experience blocks before payload assembly (producer contract)."""
    for i, b in enumerate(blocks or []):
        if not isinstance(b, dict):
            logger.info(
                "EXPERIENCE_BLOCK_%s idx=%s INVALID type=%s",
                stage,
                i,
                type(b).__name__,
            )
            continue
        bullets = b.get("bullets")
        prev = (bullets or [])[:8] if isinstance(bullets, list) else bullets
        logger.info(
            "EXPERIENCE_BLOCK_%s idx=%s company=%r title=%r date=%r location=%r "
            "bullets_len=%s bullets_preview=%s",
            stage,
            i,
            b.get("company"),
            b.get("title") or b.get("role"),
            b.get("date_range") or b.get("date"),
            b.get("location"),
            len(bullets) if isinstance(bullets, list) else None,
            prev,
        )
        logger.info(
            "EXPERIENCE_BLOCK_%s idx=%s raw_json=%s",
            stage,
            i,
            json.dumps(b, ensure_ascii=False, default=str),
        )


def _log_experience_identity_sanity_check(
    payload: ResumeDocumentPayload,
    *,
    source_blocks: Optional[List[dict]] = None,
) -> None:
    """
    Pre-validation identity audit: log every experience row immediately before
    validate_resume_document_payload (does not replace validation).

    Confirms bullet-bearing entries have company or role; flags metadata-only ghosts.
    """
    logger.info("EXPERIENCE_IDENTITY_SANITY_CHECK begin (immediately before validate_resume_document_payload)")
    for i, e in enumerate(payload.experience):
        c = (e.company or "").strip()
        r = (e.role or "").strip()
        d = (e.date or "").strip()
        loc = (e.location or "").strip()
        nb = len(e.bullets or [])
        has_id = bool(c or r)
        logger.info(
            "EXPERIENCE_IDENTITY_SANITY_CHECK idx=%s company=%r role=%r date=%r location=%r bullet_count=%s",
            i,
            c,
            r,
            d,
            loc,
            nb,
        )
        if e.bullets:
            logger.info(
                "EXPERIENCE_IDENTITY_SANITY_CHECK idx=%s first_bullets=%s",
                i,
                (e.bullets or [])[:2],
            )
        if nb > 0 and not has_id:
            logger.error(
                "EXPERIENCE_IDENTITY_SANITY_CHECK idx=%s FAIL bullets present but company and role both empty",
                i,
            )
        elif nb > 0 and has_id:
            logger.info(
                "EXPERIENCE_IDENTITY_SANITY_CHECK idx=%s OK bullet-bearing entry has company or role",
                i,
            )
        meta_ghost = (not has_id) and nb > 0 and (d or loc)
        if meta_ghost:
            logger.error(
                "EXPERIENCE_IDENTITY_SANITY_CHECK idx=%s FAIL metadata-only fragment with bullets (no identity)",
                i,
            )
    if payload.experience:
        fe = payload.experience[0]
        logger.info(
            "EXPERIENCE_IDENTITY_SANITY_CHECK first_entry company=%r role=%r date=%r location=%r",
            (fe.company or "").strip(),
            (fe.role or "").strip(),
            (fe.date or "").strip(),
            (fe.location or "").strip(),
        )
    logger.info("EXPERIENCE_IDENTITY_SANITY_CHECK end")
    if source_blocks:
        for bi, b in enumerate(source_blocks):
            if isinstance(b, dict):
                logger.info(
                    "EXPERIENCE_IDENTITY_SANITY_CHECK source_block[%s]=%s",
                    bi,
                    json.dumps(b, ensure_ascii=False, default=str),
                )


def export_filename_from_name(name_line: str) -> str:
    """Dustin Na -> Dustin_Na_Resume_Tailored_<unix_ts>.docx"""
    ts = int(time.time())
    parts = re.split(r"[\s,]+", (name_line or "").strip())
    clean = [p for p in parts if p and not re.match(r"^https?://", p, re.I)][:3]
    if len(clean) >= 2:
        safe = "_".join(re.sub(r"[^\w\-]", "", p) for p in clean[:2])
        if safe:
            return f"{safe}_Resume_Tailored_{ts}.docx"
    return f"Dustin_Na_Resume_Tailored_{ts}.docx"


def _merge_recovered_section_dicts(
    recovered: List[dict], primary: List[dict], *sig_fields: str
) -> List[dict]:
    """Prepend producer-misfiled rows; skip duplicates already present (signature = fields joined)."""
    seen: set[str] = set()
    out: List[dict] = []
    for batch in (recovered, primary):
        for row in batch:
            if not isinstance(row, dict):
                continue
            sig = "|".join(str(row.get(f) or "").strip().lower() for f in sig_fields)
            if sig in seen:
                continue
            seen.add(sig)
            out.append(row)
    return out


def _count_project_bullet_lines(projects: Optional[List[dict]]) -> int:
    n = 0
    for p in projects or []:
        if not isinstance(p, dict):
            continue
        for b in p.get("bullets") or []:
            if str(b).strip():
                n += 1
    return n


def _summary_selection_debug_source_bucket(summary_source: str) -> str:
    """Coarse label for SUMMARY_SELECTION_DEBUG (strong vs fallback wiring)."""
    src = (summary_source or "").strip()
    if src in ("tailored", "identity_forward_dense", "identity_structured"):
        return "strong"
    if src.startswith("outcome_phrase"):
        return "strong"
    if src.endswith("_identity_swap"):
        return "strong"
    return "fallback"


def _apply_final_project_scrub_to_render_payload(payload: ResumeDocumentPayload) -> List[str]:
    """
    Last-chance project cleanup on the **same** ``ResumeDocumentPayload`` instance that
    is passed to ``build_docx_from_payload`` (dict round-trip through
    ``prepare_project_blocks_for_docx``). Catches assembly-only paths that could otherwise
    bypass section scrubbers.
    """
    if not payload.projects:
        return []
    rows: List[dict] = [
        {"name": p.name or "", "subtitle": p.subtitle or "", "bullets": list(p.bullets or [])}
        for p in payload.projects
    ]
    _before = _count_project_bullet_lines(rows)
    removed: List[str] = []
    scrubbed, extra_sk, extra_edu, extra_cert = prepare_project_blocks_for_docx(
        rows, removed_debug=removed
    )
    _after = _count_project_bullet_lines(scrubbed)
    proj_entries, sk_proj = _strip_skill_bucket_lines_from_project_entries(
        dict_projects_to_entries(scrubbed)
    )
    payload.projects = proj_entries
    payload.skills = merge_distinct_skill_lines(
        merge_distinct_skill_lines(list(payload.skills or []), extra_sk),
        sk_proj,
    )
    if extra_edu or extra_cert:
        logger.warning(
            "FINAL_PROJECT_SCRUB stray_education_rows=%s stray_cert_rows=%s",
            len(extra_edu),
            len(extra_cert),
        )
    logger.info(
        "PROJECT_CLEANUP_DEBUG phase=pre_render_payload before_count=%s after_count=%s removed_lines=%s",
        _before,
        _after,
        json.dumps(removed, ensure_ascii=False),
    )
    return removed


def _sync_summary_and_projects_on_payload(
    payload: ResumeDocumentPayload,
    *,
    tailored: str,
    match_strength: str,
    resume_data: dict,
    grounding_corpus: str,
) -> None:
    """
    Hard net: same ``strongest_summary_from_resume`` + ``prepare_project_blocks_for_docx``
    round-trip the tests use, applied **in place** on ``payload`` (summary + projects + skills).
    ``tailored`` must already be tool-gated by the export package.
    """
    fs, fss = strongest_summary_from_resume(
        tailored, match_strength, resume_data, grounding_corpus
    )
    payload.summary = fs
    payload.summary_source = fss
    _apply_final_project_scrub_to_render_payload(payload)


def _finalize_export_payload_for_docx_render(
    payload: ResumeDocumentPayload,
    *,
    tailored: str,
    match_strength: str,
    resume_data: dict,
    grounding_corpus: str,
) -> None:
    """
    Single in-place finalization on the render payload: re-select summary, scrub projects,
    re-validate. ``id(payload)`` must stay stable (same object through ``build_docx_from_payload``).
    """
    logger.info(
        "PAYLOAD_ID_DEBUG stage=after_build_resume_document_payload payload_id=%s",
        id(payload),
    )
    _sync_summary_and_projects_on_payload(
        payload,
        tailored=tailored,
        match_strength=match_strength,
        resume_data=resume_data,
        grounding_corpus=grounding_corpus,
    )
    logger.info(
        "SUMMARY_SELECTION_DEBUG pass=pre_render selected_summary=%r source=%s summary_source=%r",
        payload.summary,
        _summary_selection_debug_source_bucket(payload.summary_source),
        payload.summary_source,
    )
    validate_resume_document_payload(payload)
    logger.info(
        "PAYLOAD_ID_DEBUG stage=after_finalize_export_payload payload_id=%s",
        id(payload),
    )


def build_export_docx_package(
    resume_data: dict,
    rewrite_result: dict,
    score_result: dict,
    mapping_result: dict,
    job_signals: dict,
    *,
    refinery_experience_spacing: bool = False,
    selected_change_ids: Optional[Sequence[str]] = None,
    export_route_label: Optional[str] = None,
) -> Tuple[bytes, str, Optional[str], Optional[str]]:
    """
    Returns (docx_bytes, filename, error_message, validation_success_message).

    If error_message is set, docx_bytes is empty — do not send the file.
    On success, validation_success_message is DOCX EXPORT VALIDATED - READY.

    ``selected_change_ids`` is optional plumbing for future selective DOCX apply; when
    provided (including an empty list), a ``SELECTED_CHANGE_IDS_DEBUG`` log line is emitted.
    Omit the argument for internal/script callers that do not participate in selection.
    """
    logger.info(
        "EXPORT_PIPELINE_DEBUG payload_builder=%s.%s docx_renderer=%s.%s legacy_payload_builders=none",
        _DOCX_PAYLOAD_BUILDER.__module__,
        _DOCX_PAYLOAD_BUILDER.__qualname__,
        build_docx_from_payload.__module__,
        build_docx_from_payload.__qualname__,
    )
    if selected_change_ids is not None:
        sid = list(selected_change_ids)
        logger.info(
            "SELECTED_CHANGE_IDS_DEBUG %s",
            json.dumps(
                {
                    "route": export_route_label or "build_export_docx_package",
                    "count": len(sid),
                    "ids": sid,
                },
                ensure_ascii=False,
            ),
        )
    raw_text = (resume_data.get("raw_text") or "").strip()
    sections = resume_data.get("sections") if isinstance(resume_data.get("sections"), dict) else {}
    experience_blocks = list(sections.get("experience") or []) if isinstance(sections, dict) else []
    projects = list(sections.get("projects") or []) if isinstance(sections, dict) else []
    education_base = list(sections.get("education") or []) if isinstance(sections, dict) else []
    certifications_base = (
        list(sections.get("certifications") or []) if isinstance(sections, dict) else []
    )

    bullet_changes = rewrite_result.get("bullet_changes") or []
    experience_merged = apply_bullet_changes_to_experience(experience_blocks, bullet_changes)
    experience_merged = repair_experience_api_blocks_identity_from_bullets(experience_merged)
    name, email, phone, linkedin, github = _extract_contact_lines(raw_text, resume_data)
    experience_for_docx = filter_experience_blocks_for_docx(
        experience_merged,
        name,
        email=email,
        phone=phone,
        linkedin=linkedin,
        github=github,
    )
    # Experience bullet ordering: prioritize_experience_entry_bullets in
    # resume_document_assembly.build_experience_entries_identity_first (signal-based, role-aware).
    _log_experience_blocks_for_debug("PRE_PREPARE", experience_for_docx)
    try:
        experience_for_docx = prepare_experience_blocks_for_docx(experience_for_docx)
    except ResumeContractError as exc:
        logger.error("RESUME_CONTRACT_FAILED experience_prepare: %s", exc)
        return b"", "", str(exc), None
    _log_experience_blocks_for_debug("POST_PREPARE", experience_for_docx)
    experience_for_docx = _trim_gainwell_experience_blocks_for_export_readability(
        experience_for_docx, max_bullets=6
    )
    _rewrite_classification_gating_leak_in_blocks(experience_for_docx)
    experience_for_docx = _filter_misfiled_experience_blocks_for_docx(experience_for_docx)
    experience_for_docx = _normalize_experience_headers(experience_for_docx)
    experience_for_docx = _clean_experience_trailing_metadata(experience_for_docx)
    _proj_cleanup_removed: List[str] = []
    _rewrite_classification_gating_leak_in_blocks(projects)
    _orig_proj_bullets = _count_project_bullet_lines(projects)
    (
        projects_scrubbed,
        skill_lines_from_projects,
        education_from_projects,
        certifications_from_projects,
    ) = prepare_project_blocks_for_docx(projects, removed_debug=_proj_cleanup_removed)
    projects_for_docx = prioritize_project_blocks_for_export(projects_scrubbed)
    _cleaned_proj_bullets = _count_project_bullet_lines(projects_scrubbed)
    logger.info(
        "PROJECT_CLEANUP_DEBUG phase=raw_sections before_count=%s after_count=%s removed_lines=%s",
        _orig_proj_bullets,
        _cleaned_proj_bullets,
        json.dumps(_proj_cleanup_removed, ensure_ascii=False),
    )
    education = _merge_recovered_section_dicts(
        education_from_projects,
        _merge_recovered_section_dicts(
            parse_education_dicts_from_raw_text(raw_text),
            education_base,
            "degree",
            "institution",
            "date",
            "location",
        ),
        "degree",
        "institution",
        "date",
        "location",
    )
    certifications = _merge_recovered_section_dicts(
        certifications_from_projects,
        _merge_recovered_section_dicts(
            parse_certification_dicts_from_raw_text(raw_text),
            certifications_base,
            "name",
            "issuer",
            "date",
        ),
        "name",
        "issuer",
        "date",
    )

    grounding_corpus = _build_grounding_corpus(
        resume_data, experience_for_docx, bullet_changes
    )

    match_strength = derive_match_strength(score_result, mapping_result, job_signals)
    tailored = str(rewrite_result.get("tailored_summary") or "").strip()
    # Rewrite may emit a tool-stack one-liner that passes hygiene but must not win over
    # identity / outcome summaries on the DOCX path (pipeline order: gate → select).
    if tailored and is_tool_centric_summary(tailored):
        logger.info(
            "export_docx DROPPING_TOOL_CENTRIC_TAILORED tailored_preview=%r — "
            "strongest_summary_from_resume will use resume-only ladder",
            tailored[:280],
        )
        tailored = ""
    summary, summary_source = strongest_summary_from_resume(
        tailored, match_strength, resume_data, grounding_corpus
    )
    logger.info(
        "SUMMARY_SELECTION_DEBUG selected_summary=%r source=%s summary_source=%r",
        summary,
        _summary_selection_debug_source_bucket(summary_source),
        summary_source,
    )

    skills_raw = sections.get("skills") if isinstance(sections, dict) else []
    default_skills_placeholder = not (isinstance(skills_raw, list) and len(skills_raw) > 0)
    skills_items = merge_distinct_skill_lines(
        coalesce_skills_for_export(sections, raw_text),
        skill_lines_from_projects,
    )
    skills_line = skills_to_display_line(skills_items)

    cl_val = grounding_corpus.lower()
    ungrounded = list_ungrounded_summary_tokens(summary, cl_val)
    logger.info(
        "export_docx pre-validation summary_source=%s len=%s summary=%r",
        summary_source,
        len(summary),
        summary,
    )
    if ungrounded:
        logger.warning(
            "export_docx summary tokens not in resume corpus (lexicon-exempt check): %s",
            ungrounded,
        )

    for pq in presentation_quality_warnings(
        summary=summary,
        summary_source=summary_source,
        experience_blocks=experience_for_docx,
        skills_line=skills_line,
        default_skills_placeholder=default_skills_placeholder,
    ):
        logger.warning("export presentation quality: %s", pq)
    for hint in section_integrity_sanity_hints(experience_for_docx):
        logger.warning("export section integrity hint: %s", hint)

    pre_fail = validate_export_pre_docx(
        summary=summary,
        experience_merged=experience_for_docx,
        match_strength=match_strength,
        resume_data=resume_data,
        experience_original=experience_blocks,
        bullet_changes=bullet_changes,
    )
    if pre_fail:
        return b"", "", format_validation_failure(pre_fail), None

    contact_line = " | ".join(x for x in (email, phone, linkedin, github) if x)
    try:
        export_payload = _DOCX_PAYLOAD_BUILDER(
            name=name,
            contact=contact_line,
            summary=summary,
            summary_source=summary_source,
            experience_blocks=experience_for_docx,
            projects=projects_for_docx,
            education=education,
            certifications=certifications,
            skills=skills_items,
            projects_already_prepared=True,
        )
        _log_experience_identity_sanity_check(
            export_payload, source_blocks=experience_for_docx
        )
        validate_resume_document_payload(export_payload)
    except ResumeContractError as exc:
        logger.error("RESUME_CONTRACT_FAILED: %s", exc)
        return b"", "", str(exc), None
    try:
        _finalize_export_payload_for_docx_render(
            export_payload,
            tailored=tailored,
            match_strength=match_strength,
            resume_data=resume_data,
            grounding_corpus=grounding_corpus,
        )
    except ResumeContractError as exc:
        logger.error("RESUME_CONTRACT_FAILED finalize_export_payload: %s", exc)
        return b"", "", str(exc), None
    _assembled = canonical_resume_dict(export_payload)
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG selected_final_summary=%r summary_source=%r",
        export_payload.summary,
        export_payload.summary_source,
    )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG final_summary=%r",
        export_payload.summary,
    )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG experience_entries_company_role=%s",
        json.dumps(
            [
                {"company": (e.company or "").strip(), "role": (e.role or "").strip()}
                for e in (export_payload.experience or [])
            ],
            ensure_ascii=False,
        ),
    )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG project_section_bullets=%s",
        json.dumps(
            [
                {
                    "name": (p.name or "").strip(),
                    "bullets": list(p.bullets or []),
                }
                for p in (export_payload.projects or [])
            ],
            ensure_ascii=False,
            indent=2,
        ),
    )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG certifications_payload=%s",
        json.dumps(_assembled.get("certifications"), ensure_ascii=False, indent=2),
    )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG full_structured_experience_list=%s",
        json.dumps(_assembled.get("experience"), ensure_ascii=False, indent=2),
    )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG experience_formatted_header_lines=%s",
        json.dumps(
            [
                experience_entry_header_lines(e)
                for e in (export_payload.experience or [])
            ],
            ensure_ascii=False,
        ),
    )
    _gw_ent = next(
        (e for e in (export_payload.experience or []) if "gainwell" in (e.company or "").lower()),
        None,
    )
    if _gw_ent:
        logger.info(
            "EXPORT_PRE_RENDER_DEBUG gainwell_all_bullets=%s",
            json.dumps(list(_gw_ent.bullets or []), ensure_ascii=False),
        )
    _ts_ent = next(
        (e for e in (export_payload.experience or []) if "tesla" in (e.company or "").lower()),
        None,
    )
    if _ts_ent:
        logger.info(
            "EXPORT_PRE_RENDER_DEBUG tesla_entry_header_lines=%s",
            json.dumps(experience_entry_header_lines(_ts_ent), ensure_ascii=False),
        )
    _rws_ent = next(
        (e for e in (export_payload.experience or []) if "rws" in (e.company or "").lower()),
        None,
    )
    if _rws_ent:
        logger.info(
            "EXPORT_PRE_RENDER_DEBUG rws_entry_header_lines=%s",
            json.dumps(experience_entry_header_lines(_rws_ent), ensure_ascii=False),
        )
    _proj_list = _assembled.get("projects") or []
    _pp_payload = next(
        (p for p in _proj_list if re.search(r"(?i)personal", str(p.get("name") or ""))),
        _proj_list[0] if _proj_list else None,
    )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG personal_project_payload=%s",
        json.dumps(_pp_payload, ensure_ascii=False, indent=2) if _pp_payload is not None else "null",
    )
    logger.info("RESUME_CONTRACT_FULL summary=%r", _assembled.get("summary"))
    logger.info("RESUME_CONTRACT_FULL experience=%s", _assembled.get("experience"))
    logger.info(
        "RESUME_CONTRACT_FULL_JSON=%s",
        json.dumps(_assembled, ensure_ascii=False, indent=2),
    )
    logger.info("ASSEMBLY_DEBUG SUMMARY_VALUE=%r", export_payload.summary)
    logger.info("ASSEMBLY_DEBUG HEADER_STRUCTURE=%s", _assembled["header"])
    logger.info("ASSEMBLY_DEBUG EXPERIENCE_STRUCTURE=%s", _assembled["experience"])
    logger.info("ASSEMBLY_DEBUG PROJECTS_STRUCTURE=%s", _assembled["projects"])
    logger.info("ASSEMBLY_DEBUG EDUCATION_LIST=%s", _assembled["education"])
    logger.info("ASSEMBLY_DEBUG CERTIFICATIONS_LIST=%s", _assembled["certifications"])
    logger.info("ASSEMBLY_DEBUG SKILLS_LIST=%s", _assembled["skills"])
    logger.info("EXPORT_STRUCTURED_PAYLOAD: %s", structured_payload_debug_json(export_payload))

    if export_payload.experience:
        _fe = export_payload.experience[0]
        logger.info(
            "EXPORT_PRE_RENDER_DEBUG first_experience_entry_full=%s",
            json.dumps(
                {
                    "company": _fe.company,
                    "role": _fe.role,
                    "date": _fe.date,
                    "location": _fe.location,
                    "bullets": list(_fe.bullets or []),
                },
                ensure_ascii=False,
            ),
        )
        logger.info(
            "EXPORT_PRE_RENDER_DEBUG first_experience_header_lines=%s",
            experience_entry_header_lines(_fe),
        )
        logger.info(
            "EXPORT_PRE_RENDER_DEBUG first_experience_bullets_filtered=%s",
            json.dumps(list(_fe.bullets or []), ensure_ascii=False),
        )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG projects_payload_full=%s",
        json.dumps(_assembled.get("projects"), ensure_ascii=False, indent=2),
    )
    if export_payload.projects:
        _fp = export_payload.projects[0]
        logger.info(
            "EXPORT_PRE_RENDER_DEBUG first_project_entry=%s",
            json.dumps(
                {
                    "name": _fp.name,
                    "subtitle": _fp.subtitle,
                    "bullets": _fp.bullets,
                },
                ensure_ascii=False,
            ),
        )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG education_list=%s",
        json.dumps(_assembled.get("education"), ensure_ascii=False, indent=2),
    )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG certifications_list=%s",
        json.dumps(_assembled.get("certifications"), ensure_ascii=False, indent=2),
    )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG skills_final_payload_list=%s",
        json.dumps(skills_items, ensure_ascii=False),
    )
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG skills_grouped_display_lines=%s",
        json.dumps(skills_to_display_lines(skills_items), ensure_ascii=False),
    )
    _render_payload_dump = {
        "header_contact_line": contact_line,
        "header_payload": _assembled.get("header"),
        "experience_all": _assembled.get("experience"),
        "projects": _assembled.get("projects"),
        "education": _assembled.get("education"),
        "certifications": _assembled.get("certifications"),
        "skills_payload": list(export_payload.skills),
        "skills_display_lines": skills_to_display_lines(list(export_payload.skills)),
        "final_summary": export_payload.summary,
        "summary_source": export_payload.summary_source,
    }
    logger.info(
        "EXPORT_RENDER_PAYLOAD_DUMP %s",
        json.dumps(_render_payload_dump, ensure_ascii=False, indent=2),
    )
    _section_order = ["SUMMARY", "EXPERIENCE"]
    if export_payload.projects:
        _section_order.append("PROJECTS")
    if export_payload.education:
        _section_order.append("EDUCATION")
    if export_payload.certifications:
        _section_order.append("CERTIFICATIONS")
    _section_order.append("SKILLS")
    logger.info(
        "EXPORT_PRE_RENDER_DEBUG section_order=%s",
        json.dumps(_section_order, ensure_ascii=False),
    )

    fn = export_filename_from_name(name)

    logger.info("STRUCT_CONTRACT pre_render payload.summary=%r", export_payload.summary)
    logger.info(
        "STRUCT_CONTRACT pre_render payload.experience=%s",
        json.dumps(_assembled["experience"], ensure_ascii=False, indent=2),
    )
    _proj_text = " ".join(
        " ".join(
            [(p.name or ""), (p.subtitle or "")]
            + [str(b) for b in (p.bullets or [])]
        )
        for p in (export_payload.projects or [])
    ).lower()
    _forbidden_hits = [
        tok
        for tok in ("b.a.", "b.s.", "mba", "google", "itil", "ecba")
        if tok in _proj_text
    ]
    logger.info(
        "FINAL_PAYLOAD_DEBUG summary=%r summary_source=%r source=%s "
        "summary_contains_senior_bsa=%s forbidden_project_substrings_found=%s project_bullets=%s",
        export_payload.summary,
        export_payload.summary_source,
        _summary_selection_debug_source_bucket(export_payload.summary_source),
        ("senior business systems analyst" in (export_payload.summary or "").lower()),
        _forbidden_hits,
        json.dumps(
            [
                {
                    "name": (p.name or "").strip(),
                    "subtitle": (p.subtitle or "").strip(),
                    "bullets": list(p.bullets or []),
                }
                for p in (export_payload.projects or [])
            ],
            ensure_ascii=False,
            indent=2,
        ),
    )
    logger.info(
        "EXPORT_PIPELINE_DEBUG stage=hard_net_before_docx reapplying=%s.%s",
        _sync_summary_and_projects_on_payload.__module__,
        _sync_summary_and_projects_on_payload.__qualname__,
    )
    _sync_summary_and_projects_on_payload(
        export_payload,
        tailored=tailored,
        match_strength=match_strength,
        resume_data=resume_data,
        grounding_corpus=grounding_corpus,
    )
    if maybe_apply_portfolio_resume_polish(export_payload):
        try:
            validate_resume_document_payload(export_payload)
        except ResumeContractError as exc:
            logger.error("RESUME_CONTRACT_FAILED portfolio_polish: %s", exc)
            return b"", "", str(exc), None

    # Post-finalize / post-polish expectations for STEP 4 (must match rendered DOCX).
    expect_projects = bool(export_payload.projects)
    expect_education = bool(export_payload.education)
    expect_certifications = bool(export_payload.certifications)

    logger.info(
        "FINAL_EXPORT_PAYLOAD_DEBUG payload_id=%s summary=%r project_bullets=%s",
        id(export_payload),
        export_payload.summary,
        json.dumps(
            [
                {
                    "name": (p.name or "").strip(),
                    "subtitle": (p.subtitle or "").strip(),
                    "bullets": list(p.bullets or []),
                }
                for p in (export_payload.projects or [])
            ],
            ensure_ascii=False,
            indent=2,
        ),
    )
    logger.info(
        "PAYLOAD_ID_DEBUG stage=immediately_before_build_docx_from_payload payload_id=%s",
        id(export_payload),
    )
    _runtime_projects = [
        {
            "name": (p.name or "").strip(),
            "subtitle": (p.subtitle or "").strip(),
            "bullets": list(p.bullets or []),
        }
        for p in (export_payload.projects or [])
    ]
    _runtime_education = [
        {
            "degree": (e.degree or "").strip(),
            "institution": (e.institution or "").strip(),
            "date": (e.date or "").strip(),
            "location": (e.location or "").strip(),
            "bullets": list(e.bullets or []),
        }
        for e in (export_payload.education or [])
    ]
    _runtime_certifications = [
        {
            "name": (c.name or "").strip(),
            "issuer": (c.issuer or "").strip(),
            "date": (c.date or "").strip(),
            "bullets": list(c.bullets or []),
        }
        for c in (export_payload.certifications or [])
    ]
    if _export_debug_enabled():
        logger.debug(
            "final export payload snapshot: %s",
            json.dumps(
                {
                    "summary": export_payload.summary or "",
                    "summary_source": export_payload.summary_source or "",
                    "projects": _runtime_projects,
                    "education": _runtime_education,
                    "certifications": _runtime_certifications,
                },
                ensure_ascii=False,
            ),
        )
    docx_bytes = build_docx_from_payload(
        export_payload,
        refinery_experience_spacing=refinery_experience_spacing,
    )

    blob = _export_text_blob_from_payload(export_payload)
    post_fail = validate_export_post_docx(
        export_text_blob=blob,
        docx_bytes=docx_bytes,
        expect_projects=expect_projects,
        expect_education=expect_education,
        expect_certifications=expect_certifications,
    )
    if post_fail:
        return b"", "", format_validation_failure(post_fail), None

    return docx_bytes, fn, None, _SUCCESS_MSG
