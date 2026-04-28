"""DOCX experience section: identity lines must precede bullets; bullets use List Bullet style."""

from __future__ import annotations

import unittest
from io import BytesIO

from docx import Document

from app.services.export_docx import build_docx_from_payload
from app.services.resume_document_assembly import build_resume_document_payload

_LIST_BULLET = "List Bullet"
_NORMAL = "Normal"


def _experience_section_meta(docx_bytes: bytes) -> list[tuple[str, str]]:
    """(text, style_name) for non-empty paragraphs between EXPERIENCE and the next section."""
    doc = Document(BytesIO(docx_bytes))
    in_exp = False
    out: list[tuple[str, str]] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if t.upper() == "EXPERIENCE":
            in_exp = True
            continue
        if in_exp and t.upper() in ("PROJECTS", "EDUCATION", "CERTIFICATIONS", "SKILLS"):
            break
        if in_exp:
            out.append((t, p.style.name))
    return out


def _split_experience_entries(meta: list[tuple[str, str]]) -> list[list[tuple[str, str]]]:
    """New job starts when a non–List-Bullet paragraph follows a List Bullet paragraph."""
    if not meta:
        return []
    starts: list[int] = [0]
    for i in range(1, len(meta)):
        prev_b = meta[i - 1][1] == _LIST_BULLET
        cur_b = meta[i][1] == _LIST_BULLET
        if prev_b and not cur_b:
            starts.append(i)
    out: list[list[tuple[str, str]]] = []
    for si, s in enumerate(starts):
        end = starts[si + 1] if si + 1 < len(starts) else len(meta)
        out.append(meta[s:end])
    return out


def _assert_entry_identity_then_bullets_styled(
    segment: list[tuple[str, str]], *, msg: str
) -> None:
    styles = [s for _, s in segment]
    texts = [t for t, _ in segment]
    if not segment:
        raise AssertionError(f"empty entry segment: {msg}")
    if styles[0] == _LIST_BULLET:
        raise AssertionError(f"bullet before identity: first line={texts[0]!r} ({msg})")
    try:
        ib = styles.index(_LIST_BULLET)
    except ValueError as exc:
        raise AssertionError(f"no List Bullet paragraphs in entry ({msg})") from exc
    for j in range(ib):
        if styles[j] == _LIST_BULLET:
            raise AssertionError(f"bullet in identity block at {j}: {texts[j]!r} ({msg})")
    for j in range(ib, len(styles)):
        if styles[j] != _LIST_BULLET:
            raise AssertionError(
                f"non–List-Bullet after bullets started: {texts[j]!r} style={styles[j]!r} ({msg})"
            )


def _assert_entries_visually_separated_styled(segments: list[list[tuple[str, str]]]) -> None:
    if len(segments) < 2:
        return
    for a, b in zip(segments, segments[1:]):
        if not a or not b:
            continue
        if a[-1][1] == _LIST_BULLET and b[0][1] == _LIST_BULLET:
            raise AssertionError("two entries glued without a new identity header")


class TestDocxExperienceRenderOrder(unittest.TestCase):
    def test_company_role_date_before_bullets_per_entry(self) -> None:
        payload = build_resume_document_payload(
            name="Pat Lee",
            contact="pat@example.com",
            summary="Business analyst with delivery and documentation experience.",
            summary_source="test",
            experience_blocks=[
                {
                    "company": "Alpha Corp",
                    "title": "Senior Analyst",
                    "date_range": "2022 – Present",
                    "location": "Remote",
                    "bullets": ["First bullet for Alpha.", "Second bullet for Alpha."],
                },
                {
                    "company": "Beta LLC",
                    "title": "Analyst",
                    "date_range": "2020 – 2022",
                    "location": "Austin, TX",
                    "bullets": ["Single bullet for Beta."],
                },
            ],
            projects=[],
            education=[],
            certifications=[],
            skills=["SQL"],
        )
        docx_bytes = build_docx_from_payload(payload)
        body = _experience_section_meta(docx_bytes)

        self.assertGreaterEqual(len(body), 10, msg=body)
        self.assertEqual(body[0], ("Alpha Corp", _NORMAL))
        self.assertEqual(body[1], ("Senior Analyst", _NORMAL))
        self.assertIn("Remote", body[2][0])
        self.assertEqual(body[2][1], _NORMAL)
        self.assertIn("2022", body[3][0])
        self.assertEqual(body[3][1], _NORMAL)
        self.assertEqual(body[4][1], _LIST_BULLET)
        self.assertIn("Alpha", body[4][0])
        self.assertEqual(body[5][1], _LIST_BULLET)
        self.assertEqual(body[6], ("Beta LLC", _NORMAL))
        self.assertEqual(body[7], ("Analyst", _NORMAL))
        self.assertIn("Austin", body[8][0])
        self.assertEqual(body[8][1], _NORMAL)
        self.assertIn("2020", body[9][0])
        self.assertEqual(body[9][1], _NORMAL)
        self.assertEqual(body[10][1], _LIST_BULLET)

        segs = _split_experience_entries(body)
        self.assertEqual(len(segs), 2)
        for i, seg in enumerate(segs):
            _assert_entry_identity_then_bullets_styled(seg, msg=f"alpha_beta[{i}]")
        _assert_entries_visually_separated_styled(segs)

    def test_gainwell_first_three_jobs_rendering_sanity(self) -> None:
        """
        Rendering sanity: identity-first order, List Bullet only for bullets, Gainwell first
        with Senior BSA / UAT Lead role.
        """
        payload = build_resume_document_payload(
            name="Jane Doe",
            contact="jane@example.com",
            summary="Business analyst with delivery documentation and validation experience.",
            summary_source="test",
            experience_blocks=[
                {
                    "company": "Gainwell Technologies",
                    "title": "Senior Business Systems Analyst / Senior UAT Lead",
                    "date_range": "April 2024 – Present",
                    "location": "Remote",
                    "bullets": ["Led UAT for Provider Portal."],
                },
                {
                    "company": "Tesla",
                    "title": "Data Specialist (Autopilot)",
                    "date_range": "July 2020 – June 2022",
                    "location": "San Mateo, CA",
                    "bullets": ["Supported Autopilot data pipelines."],
                },
                {
                    "company": "RWS Moravia (Client: Apple)",
                    "title": "Business Data Technician",
                    "date_range": "June 2019 – June 2020",
                    "location": "Sunnyvale, CA",
                    "bullets": ["Maintained localization data workflows."],
                },
            ],
            projects=[],
            education=[],
            certifications=[],
            skills=["SQL"],
        )
        docx_bytes = build_docx_from_payload(payload)
        body = _experience_section_meta(docx_bytes)

        self.assertGreater(len(body), 0, msg=body)
        self.assertIn("Gainwell", body[0][0], msg=f"Gainwell must lead experience: {body[:5]!r}")
        self.assertEqual(body[0][1], _NORMAL)
        self.assertIn("Senior Business Systems Analyst", body[1][0])
        self.assertIn("UAT Lead", body[1][0])
        self.assertEqual(body[1][1], _NORMAL)
        self.assertIn("Remote", body[2][0])
        self.assertEqual(body[2][1], _NORMAL)
        self.assertRegex(body[3][0], r"(2024|Present)", msg=body[3][0])
        self.assertEqual(body[3][1], _NORMAL)
        self.assertEqual(body[4][1], _LIST_BULLET)

        segments = _split_experience_entries(body)
        self.assertEqual(len(segments), 3, msg=segments)
        for k, seg in enumerate(segments):
            _assert_entry_identity_then_bullets_styled(seg, msg=f"entry[{k}]")
        _assert_entries_visually_separated_styled(segments)

        self.assertIn("Tesla", segments[1][0][0])
        self.assertIn("Autopilot", segments[1][1][0])
        self.assertEqual(segments[1][-1][1], _LIST_BULLET)
        self.assertIn("RWS", segments[2][0][0])
        self.assertIn("Business Data Technician", segments[2][1][0])


class TestDocxEducationCertPlainParagraphs(unittest.TestCase):
    def test_education_and_cert_lines_are_not_list_bullet(self) -> None:
        payload = build_resume_document_payload(
            name="Alex Kim",
            contact="a@example.com",
            summary="Professional with documented delivery experience.",
            summary_source="test",
            experience_blocks=[
                {
                    "company": "Co",
                    "title": "Role",
                    "date_range": "2020 – Present",
                    "location": "Remote",
                    "bullets": ["Did work."],
                }
            ],
            projects=[],
            education=[
                {
                    "degree": "B.S. Mathematics",
                    "institution": "State U",
                    "date": "2012",
                    "location": "",
                    "bullets": ["Honors: cum laude"],
                }
            ],
            certifications=[
                {
                    "name": "PMP",
                    "issuer": "PMI",
                    "date": "2015",
                    "bullets": ["Credential active."],
                }
            ],
            skills=["SQL"],
        )
        docx_bytes = build_docx_from_payload(payload)
        doc = Document(BytesIO(docx_bytes))
        in_edu = False
        in_cert = False
        edu_styles: list[str] = []
        cert_styles: list[str] = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if not t:
                continue
            u = t.upper()
            if u == "EDUCATION":
                in_edu = True
                in_cert = False
                continue
            if u == "CERTIFICATIONS":
                in_cert = True
                in_edu = False
                continue
            if u == "SKILLS":
                break
            if in_edu and u != "EDUCATION":
                edu_styles.append(p.style.name)
            if in_cert and u != "CERTIFICATIONS":
                cert_styles.append(p.style.name)
        self.assertTrue(edu_styles)
        self.assertTrue(cert_styles)
        self.assertNotIn(_LIST_BULLET, edu_styles, msg=edu_styles)
        self.assertNotIn(_LIST_BULLET, cert_styles, msg=cert_styles)


def _skills_section_meta(docx_bytes: bytes) -> list[tuple[str, str]]:
    """(text, style_name) for non-empty paragraphs after SKILLS heading."""
    doc = Document(BytesIO(docx_bytes))
    in_sk = False
    out: list[tuple[str, str]] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if t.upper() == "SKILLS":
            in_sk = True
            continue
        if in_sk:
            out.append((t, p.style.name))
    return out


def _projects_section_meta(docx_bytes: bytes) -> list[tuple[str, str]]:
    """(text, style_name) between PROJECTS and EDUCATION (or CERTIFICATIONS / SKILLS)."""
    doc = Document(BytesIO(docx_bytes))
    in_p = False
    out: list[tuple[str, str]] = []
    stops = frozenset(("EDUCATION", "CERTIFICATIONS", "SKILLS"))
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if t.upper() == "PROJECTS":
            in_p = True
            continue
        if in_p and t.upper() in stops:
            break
        if in_p:
            out.append((t, p.style.name))
    return out


class TestDocxRenderingTypeSanity(unittest.TestCase):
    """
    Rendering type sanity: only true bullet content uses List Bullet;
    identity and section metadata stay Normal.
    """

    def test_full_hierarchy_experience_education_cert_skills_projects(self) -> None:
        payload = build_resume_document_payload(
            name="Riley Chen",
            contact="r@example.com",
            summary="Analyst with delivery documentation and stakeholder engagement experience.",
            summary_source="test",
            experience_blocks=[
                {
                    "company": "Northwind",
                    "title": "Business Analyst",
                    "date_range": "2021 – Present",
                    "location": "Remote",
                    "bullets": [
                        "Facilitated requirements workshops with product and engineering.",
                        "Documented process flows for order-to-cash redesign.",
                    ],
                },
                {
                    "company": "Contoso",
                    "title": "Junior Analyst",
                    "date_range": "2019 – 2021",
                    "location": "Chicago, IL",
                    "bullets": ["Maintained requirements traceability in Azure DevOps."],
                },
            ],
            projects=[
                {
                    "name": "Portfolio — Requirements Dashboard",
                    "subtitle": "Personal exploration",
                    "bullets": [
                        "Built a small React UI for backlog prioritization views.",
                    ],
                }
            ],
            education=[
                {
                    "degree": "B.A. Economics",
                    "institution": "Midwest University",
                    "date": "2018",
                    "location": "",
                    "bullets": ["Dean's List"],
                }
            ],
            certifications=[
                {
                    "name": "Certified ScrumMaster",
                    "issuer": "Scrum Alliance",
                    "date": "2020",
                    "bullets": ["Renewed through 2025."],
                }
            ],
            skills=[
                "Data & Analytics: SQL, Excel",
                "Systems & Platforms: Jira, Confluence",
            ],
        )
        docx_bytes = build_docx_from_payload(payload)

        # Checks 1–4: experience — identity lines never List Bullet; bullets only List Bullet.
        exp_meta = _experience_section_meta(docx_bytes)
        self.assertGreaterEqual(len(exp_meta), 8, msg=exp_meta)
        for seg in _split_experience_entries(exp_meta):
            _assert_entry_identity_then_bullets_styled(seg, msg="sanity_experience")
            for text, st in seg:
                if st != _LIST_BULLET:
                    self.assertNotRegex(
                        text,
                        r"(?i)^\s*(facilitated|documented|maintained)\b",
                        msg=f"action-shaped line must not appear as non-bullet: {text!r}",
                    )
                else:
                    self.assertGreaterEqual(
                        len(text),
                        12,
                        msg="bullet lines should be substantive action statements",
                    )

        # Checks 5–6: education and certifications — no List Bullet in those sections.
        doc = Document(BytesIO(docx_bytes))
        section_styles: dict[str, list[str]] = {"EDUCATION": [], "CERTIFICATIONS": []}
        current: str | None = None
        for p in doc.paragraphs:
            t = p.text.strip()
            if not t:
                continue
            u = t.upper()
            if u in section_styles:
                current = u
                continue
            if u in ("SKILLS", "SUMMARY", "EXPERIENCE", "PROJECTS"):
                if u == "SKILLS":
                    break
                if u in ("SUMMARY", "EXPERIENCE", "PROJECTS"):
                    current = None
                continue
            if current in section_styles:
                section_styles[current].append(p.style.name)
        self.assertNotIn(_LIST_BULLET, section_styles["EDUCATION"], msg=section_styles)
        self.assertNotIn(_LIST_BULLET, section_styles["CERTIFICATIONS"], msg=section_styles)

        # Check 7: skills — grouped body lines, Normal (not list bullets).
        sk_meta = _skills_section_meta(docx_bytes)
        self.assertGreaterEqual(len(sk_meta), 2, msg=sk_meta)
        for text, st in sk_meta:
            self.assertEqual(st, _NORMAL, msg=f"skills line must be Normal: {text!r} {st!r}")
            self.assertIn(":", text, msg=f"expected grouped skill line with colon: {text!r}")

        # Project hierarchy: title/subtitle Normal; only bullet lines use List Bullet.
        proj_meta = _projects_section_meta(docx_bytes)
        self.assertGreaterEqual(len(proj_meta), 2, msg=proj_meta)
        self.assertEqual(proj_meta[0][1], _NORMAL)
        self.assertEqual(proj_meta[1][1], _NORMAL)
        self.assertEqual(proj_meta[2][1], _LIST_BULLET)


if __name__ == "__main__":
    unittest.main()
